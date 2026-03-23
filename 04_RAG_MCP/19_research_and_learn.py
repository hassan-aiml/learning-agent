import os
import io
import requests
import chromadb
import streamlit as st
import streamlit.components.v1 as components
from dotenv import load_dotenv
from anthropic import Anthropic
from bs4 import BeautifulSoup
from sentence_transformers import SentenceTransformer

# Load API key — works both locally (.env file) and on Streamlit Cloud (st.secrets)
load_dotenv()
if "ANTHROPIC_API_KEY" not in os.environ:
    os.environ["ANTHROPIC_API_KEY"] = st.secrets.get("ANTHROPIC_API_KEY", "")

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Learn Smarter",
    page_icon="🧠",
    layout="centered"
)

# ── Global styles: tab colors + mobile viewport ───────────────────────────────
# st.html() is used here because st.markdown() can strip or misrender
# <style> blocks in some Streamlit versions, causing CSS to appear as text.
st.html("""
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>

/* ── Learn Smarter — Blue Light theme ── */
:root {
    --navy:          #1e3a5f;
    --navy-mid:      #162d4a;
    --blue:          #1d4ed8;
    --blue-mid:      #1e40af;
    --blue-light:    #eff6ff;
    --blue-pale:     #f0f7ff;
    --blue-glow:     rgba(29,78,216,0.12);
    --sky:           #60a5fa;
    --sky-light:     #bfdbfe;
    --card-white:    #ffffff;
    --border:        #dbeafe;
    --text-primary:  #1e3a5f;
    --text-secondary:#3b5f8a;
    --text-hint:     #94a3b8;
    --text-sidebar:  #e8f0fe;
}

/* ── Page background: soft ice blue ── */
html, body,
[data-testid="stAppViewContainer"],
[data-testid="stAppViewBlockContainer"],
.stApp, .main {
    background-color: var(--blue-pale) !important;
    color: var(--text-primary) !important;
}

/* ── Content card: white ── */
.block-container {
    background: var(--card-white) !important;
    padding: 1.5rem 2rem !important;
    box-shadow: 0 1px 4px rgba(29,78,216,0.08), 0 2px 10px rgba(29,78,216,0.05) !important;
}
@media (min-width: 641px) {
    .block-container {
        border-radius: 12px !important;
        margin-top: 1rem !important;
        border: 1px solid var(--border) !important;
    }
}

/* ── Body text ── */
.block-container p,
.block-container li,
.block-container span,
.block-container label,
[data-testid="stMarkdownContainer"] p {
    color: var(--text-primary) !important;
}

/* ── Page titles ── */
h1, .block-container h1 {
    color: var(--navy) !important;
    font-weight: 700 !important;
    border-bottom: 2px solid var(--sky-light) !important;
    padding-bottom: 0.5rem !important;
    margin-bottom: 0.75rem !important;
}

/* ── Subheaders ── */
h2, .block-container h2 {
    color: var(--text-primary) !important;
    font-size: 1.1rem !important;
    font-weight: 600 !important;
}
h3, .block-container h3 {
    color: var(--blue) !important;
    font-size: 1rem !important;
    font-weight: 500 !important;
}

/* ── Captions ── */
.stCaption p, small {
    color: var(--text-secondary) !important;
}

/* ── Sidebar: deep navy ── */
section[data-testid="stSidebar"],
section[data-testid="stSidebar"] > div,
section[data-testid="stSidebar"] > div > div {
    background: var(--navy) !important;
    background-color: var(--navy) !important;
    border-right: 1px solid var(--navy-mid) !important;
}
section[data-testid="stSidebar"] * {
    color: var(--text-sidebar) !important;
}
section[data-testid="stSidebar"] div[style*="color:#fbbf24"],
section[data-testid="stSidebar"] span[style*="color:#fbbf24"],
section[data-testid="stSidebar"] div[style*="color: #fbbf24"] {
    color: #fbbf24 !important;
}
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: var(--sky) !important;
}
section[data-testid="stSidebar"] .stCaption p {
    color: #93c5fd !important;
}
section[data-testid="stSidebar"] hr {
    border-color: #2a4a70 !important;
}
section[data-testid="stSidebar"] [data-testid="stAlert"] {
    background: var(--navy-mid) !important;
    border-left: 3px solid var(--sky) !important;
}
section[data-testid="stSidebar"] [data-testid="stMetricValue"] {
    color: var(--sky) !important;
    font-size: 1.4rem !important;
}

/* ── All buttons: theme blue, white text ── */
/* Default state: solid blue matching the theme accent */
.stButton > button,
[data-testid^="stBaseButton"],
button[data-testid^="stBaseButton"] {
    background-color: #1d4ed8 !important;
    background: #1d4ed8 !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    font-size: 0.875rem !important;
    box-shadow: 0 1px 3px rgba(29,78,216,0.3) !important;
    transition: all 0.2s !important;
    cursor: pointer !important;
}
/* Hover state: lighter sky blue — clear contrast with white text */
.stButton > button:hover,
[data-testid^="stBaseButton"]:hover,
button[data-testid^="stBaseButton"]:hover {
    background-color: #2563eb !important;
    background: #2563eb !important;
    color: #ffffff !important;
    box-shadow: 0 2px 6px rgba(29,78,216,0.4) !important;
    transform: translateY(-1px) !important;
}

/* ── Remove (✕) button: red on hover — clearly destructive ── */
.stButton > button[title*="Remove"]:hover,
.stButton > button[title*="remove"]:hover {
    background-color: #dc2626 !important;
    background: #dc2626 !important;
    color: #ffffff !important;
    box-shadow: 0 2px 6px rgba(220,38,38,0.4) !important;
}

/* ── Force white text on ALL button states everywhere ── */
.stButton > button *,
.stButton > button p,
.stButton > button span,
.stButton > button div,
[data-testid^="stBaseButton"] *,
[data-testid^="stBaseButton"] p,
[data-testid^="stBaseButton"] span {
    color: #ffffff !important;
}

/* ── Chat bubbles ── */
[data-testid="stChatMessageContent"] {
    border-radius: 10px !important;
    background: #f8fafc !important;
    border: 1px solid var(--border) !important;
    color: var(--text-primary) !important;
}
[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) [data-testid="stChatMessageContent"] {
    background: var(--blue-light) !important;
    border: 1px solid var(--sky-light) !important;
    border-left: 3px solid var(--blue) !important;
}

/* ── Chat avatar icons ── */
[data-testid="stChatMessageAvatarUser"] {
    background-color: #1e3a5f !important;
    color: #ffffff !important;
}
[data-testid="stChatMessageAvatarAssistant"] {
    background-color: #fbbf24 !important;
    color: #0f172a !important;
}

/* ── Chat input ── */
[data-testid="stChatInput"],
[data-testid="stChatInput"] textarea {
    background: #fff8e6 !important;
    border: 1.5px solid #fbbf24 !important;
    border-radius: 24px !important;
    color: var(--text-primary) !important;
    transition: background 0.2s, border-color 0.2s !important;
}
[data-testid="stChatInput"]:focus-within,
[data-testid="stChatInput"]:focus-within textarea {
    background: #ffffff !important;
    border-color: var(--blue) !important;
    box-shadow: 0 0 0 2px var(--blue-glow) !important;
    outline: none !important;
}

/* ── Remove purple/default focus ring on ALL inputs ── */
*:focus, *:focus-visible {
    outline-color: var(--blue) !important;
    outline-width: 2px !important;
}
textarea:focus, input:focus {
    outline: none !important;
    box-shadow: 0 0 0 2px var(--blue-glow) !important;
    border-color: var(--blue) !important;
}

/* ── File uploader drag and drop area ── */
[data-testid="stFileUploader"] {
    background: var(--navy) !important;
    border-radius: 8px !important;
}
[data-testid="stFileUploaderDropzone"] {
    background: var(--navy) !important;
    border: 2px solid var(--sky) !important;
    border-radius: 8px !important;
}
[data-testid="stFileUploaderDropzone"] * {
    color: #e8f0fe !important;
}
[data-testid="stFileUploaderDropzone"] p,
[data-testid="stFileUploaderDropzone"] span,
[data-testid="stFileUploaderDropzone"] small {
    color: #bfdbfe !important;
}
[data-testid="stFileUploaderDropzone"]:hover {
    border-color: var(--blue) !important;
    background: var(--navy-mid) !important;
}

/* ── Alerts ── */
[data-testid="stAlert"] {
    border-radius: 8px !important;
    background: var(--blue-light) !important;
    border: 1px solid var(--border) !important;
}

/* ── Sidebar info box (accepted formats) — dotted border ── */
section[data-testid="stSidebar"] [data-testid="stAlert"] {
    background: var(--navy-mid) !important;
    border: 2px dotted var(--sky) !important;
    border-radius: 8px !important;
}

/* ── Expanders ── */
[data-testid="stExpander"] {
    background: var(--card-white) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
}
[data-testid="stExpander"] summary {
    color: var(--blue) !important;
    font-weight: 500 !important;
}

/* ── Dividers ── */
hr { border-color: var(--border) !important; }

/* ── Metrics ── */
[data-testid="stMetricValue"] { color: var(--blue) !important; }
[data-testid="stMetricLabel"] { color: var(--text-secondary) !important; }

/* ── Tab bar ── */
[data-baseweb="tab-list"] {
    background: var(--card-white) !important;
    border-bottom: 2px solid var(--border) !important;
    padding: 0 4px !important;
    gap: 4px !important;
}
button[data-baseweb="tab"] {
    color: var(--text-hint) !important;
    font-weight: 500 !important;
    border-radius: 8px 8px 0 0 !important;
    padding: 10px 20px !important;
    transition: background 0.15s, color 0.15s !important;
}

/* ── Both tabs: blue highlight when selected, white text ── */
button[data-baseweb="tab"]:hover {
    color: var(--blue) !important;
    background: var(--blue-light) !important;
}
button[data-baseweb="tab"][aria-selected="true"] {
    background: var(--blue) !important;
    border-bottom: 3px solid var(--blue) !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    box-shadow: 0 -2px 8px rgba(29,78,216,0.2) !important;
}
button[data-baseweb="tab"][aria-selected="true"] * {
    color: #ffffff !important;
}

/* ── Toggle ── */
[data-testid="stToggle"] input:checked + div {
    background: var(--blue) !important;
}

/* ── Alert accents ── */
[data-testid="stAlert"][kind="success"] { border-left: 3px solid #059669 !important; }
[data-testid="stAlert"][kind="warning"]  { border-left: 3px solid #d97706 !important; }

/* ── Claude-style chat layout — full height, input pinned to bottom ── */
[data-testid="stChatInput"] {
    position: sticky !important;
    bottom: 0 !important;
    z-index: 50 !important;
    background: var(--card-white) !important;
    padding: 8px 0 4px !important;
}
/* Give chat messages area breathing room above the pinned input */
[data-testid="stChatMessage"] {
    max-width: 780px !important;
    margin: 0 auto 8px !important;
}
/* Limit message width for readability like Claude */
[data-testid="stChatMessageContent"] {
    max-width: 720px !important;
}

/* ── Hide uploaded filename under uploader — shown in Files in Chat instead ── */
[data-testid="stFileUploaderFileName"],
[data-testid="stFileUploaderFile"],
[data-testid="stFileUploader"] [data-testid="stMarkdownContainer"],
[data-testid="stFileUploader"] small,
[data-testid="stFileUploader"] .uploadedFileName {
    display: none !important;
}

/* ── Mobile ── */
@media (max-width: 640px) {

    /* Push entire app down so Streamlit top bar doesn't overlap the tabs */
    .stApp > div:first-child { padding-top: 56px !important; }
    [data-testid="stAppViewContainer"] { padding-top: 56px !important; }
    header[data-testid="stHeader"] { height: 56px !important; }

    .block-container { padding: 0.75rem 0.6rem !important; border-radius: 0 !important; border: none !important; margin-top: 0 !important; }

    /* Tabs: sticky below the top bar with a top offset matching header height */
    [data-testid="stTabs"] { margin-top: 0 !important; padding-top: 0 !important; }
    [data-baseweb="tab-list"] {
        position: sticky !important;
        top: 56px !important;
        z-index: 99 !important;
        width: 100% !important;
        margin: 0 !important;
        padding: 0 !important;
        display: flex !important;
        background: var(--card-white) !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.08) !important;
    }
    button[data-baseweb="tab"] {
        flex: 1 !important;
        min-height: 52px !important;
        font-size: 0.82rem !important;
        padding: 0 6px !important;
        justify-content: center !important;
        touch-action: manipulation !important;
    }

    h1 { font-size: 1.2rem !important; margin-top: 0.5rem !important; }
    .stButton > button { font-size: 0.85rem !important; }
    section[data-testid="stSidebar"] { min-width: 260px !important; }
}

</style>
""")


# ── Cached resources ──────────────────────────────────────────────────────────
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

@st.cache_resource
def load_claude_client():
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        st.error("ANTHROPIC_API_KEY not found. Add it to your .env file locally or to Streamlit secrets in the cloud.")
        st.stop()
    return Anthropic(api_key=api_key)

model  = load_embedding_model()
client = load_claude_client()


# ════════════════════════════════════════════════════════════════════════════════
# SHARED UTILITIES
# ════════════════════════════════════════════════════════════════════════════════

def chunk_text(text: str, chunk_size: int = 300, overlap: int = 50) -> list:
    """Split text into overlapping chunks for RAG retrieval."""
    words  = text.split()
    chunks = []
    start  = 0
    while start < len(words):
        chunk = " ".join(words[start:start + chunk_size])
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def extract_text_from_file(uploaded_file) -> str:
    """
    Extract plain text from an uploaded file.
    Supports: .txt, .md, .pdf, .docx
    """
    filename = uploaded_file.name.lower()

    if filename.endswith((".txt", ".md")):
        return uploaded_file.read().decode("utf-8")

    elif filename.endswith(".pdf"):
        try:
            import fitz
            pdf_bytes = uploaded_file.read()
            doc       = fitz.open(stream=pdf_bytes, filetype="pdf")
            return "\n\n".join([page.get_text() for page in doc])
        except ImportError:
            st.error("PyMuPDF not installed. Run: pip install pymupdf")
            return ""
        except Exception as e:
            st.error(f"Could not read PDF: {e}")
            return ""

    elif filename.endswith(".docx"):
        try:
            from docx import Document
            doc  = Document(io.BytesIO(uploaded_file.read()))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            st.error("python-docx not installed. Run: pip install python-docx")
            return ""
        except Exception as e:
            st.error(f"Could not read Word document: {e}")
            return ""

    else:
        st.error(f"Unsupported file type: {filename}")
        return ""


def build_collection_from_docs(doc_list: list, collection_name: str):
    """
    Build a ChromaDB collection from a list of
    {"name": str, "text": str} dicts.
    Returns the collection or None if empty.
    """
    all_ids   = []
    all_texts = []

    for doc in doc_list:
        chunks = chunk_text(doc["text"])
        for i, chunk in enumerate(chunks):
            all_ids.append(f"{doc['name']}_chunk_{i}")
            all_texts.append(chunk)

    if not all_texts:
        return None

    base_dir = os.path.dirname(os.path.abspath(__file__))
    chroma_client = chromadb.PersistentClient(path=os.path.join(base_dir, "knowledge_base"))
    try:
        chroma_client.delete_collection(collection_name)
    except:
        pass

    collection     = chroma_client.create_collection(collection_name)
    all_embeddings = model.encode(all_texts, show_progress_bar=False).tolist()
    collection.add(ids=all_ids, embeddings=all_embeddings, documents=all_texts)
    return collection


def retrieve_context(collection, query: str, top_k: int = 4) -> str:
    """Semantic search — returns formatted top-k chunks."""
    query_embedding = model.encode([query]).tolist()
    results = collection.query(query_embeddings=query_embedding, n_results=top_k)
    chunks  = results["documents"][0]
    return "\n\n".join([f"[Document {i+1}]: {chunk}" for i, chunk in enumerate(chunks)])


def rebuild_chat_collection():
    """
    Rebuild the active chat collection from all docs
    currently in st.session_state.chat_docs.
    Called whenever a doc is added or removed.
    """
    if not st.session_state.chat_docs:
        st.session_state.chat_collection = None
        return

    collection = build_collection_from_docs(
        st.session_state.chat_docs, "chat_docs"
    )
    st.session_state.chat_collection = collection


# ════════════════════════════════════════════════════════════════════════════════
# DEFAULT DOCUMENT — Hassan's resume from my_docs/
# ════════════════════════════════════════════════════════════════════════════════

def load_default_docs(folder: str = "my_docs") -> list:
    """
    Load all supported files from my_docs/ as a list of
    {"name": str, "text": str} dicts.
    # NOTE: The default document (Hassan's resume) is a sample to demonstrate
    # the Document Chat feature. Users can upload their own documents,
    # or use the Research Agent to find content and add it to the chat.
    """
    # Build path relative to this app file — works on Streamlit Cloud
    # where the working directory may differ from the app file location
    base_dir = os.path.dirname(os.path.abspath(__file__))
    folder   = os.path.join(base_dir, folder)

    if not os.path.exists(folder):
        return []

    docs = []
    for filename in os.listdir(folder):
        if not filename.endswith((".txt", ".md", ".pdf", ".docx")):
            continue
        filepath = os.path.join(folder, filename)
        ext      = filename.lower()
        try:
            if ext.endswith((".txt", ".md")):
                with open(filepath, "r", encoding="utf-8") as f:
                    text = f.read().strip()
            elif ext.endswith(".pdf"):
                import fitz
                doc  = fitz.open(filepath)
                text = "\n\n".join([page.get_text() for page in doc])
            elif ext.endswith(".docx"):
                from docx import Document
                doc  = Document(filepath)
                text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                continue
            if text.strip():
                docs.append({"name": filename, "text": text})
        except Exception:
            continue
    return docs


# ════════════════════════════════════════════════════════════════════════════════
# SYSTEM PROMPTS
# ════════════════════════════════════════════════════════════════════════════════

def build_chat_system_prompt(doc_names: list) -> str:
    """Dynamically build a system prompt based on active documents."""
    if not doc_names:
        return "You are a helpful assistant."

    # Special persona when only Hassan's resume is loaded
    if doc_names == ["hassan_resume.txt"] or doc_names == ["hassan_resume.md"]:
        return """You are a professional HR assistant helping recruiters and hiring
managers learn about Hassan M. Hai from his resume.
- Only answer using the resume context provided — never invent experience or skills
- If the resume doesn't contain enough information to answer, say so clearly
- Refer to the candidate by first name: Hassan
- Be objective, professional, and concise
- Do not speculate about things not mentioned in the resume"""

    names_str = ", ".join(doc_names)
    return f"""You are a helpful assistant answering questions based on the following
documents: {names_str}.

IMPORTANT RULES:
- Only answer using the document context provided — never invent information
- If the documents don't contain enough information, say so clearly
- Be concise, accurate, and helpful
- When referencing specific facts, mention which document they come from
- Do not speculate about things not mentioned in the documents"""


RESEARCH_SYSTEM_PROMPT = """You are a professional research assistant. Research topics
thoroughly and produce clear, structured reports.

Your process:
1. Search for the topic to get an overview
2. Fetch 1-2 useful URLs for more detail
3. Synthesise everything into a structured report

Report format:
# [Topic Title]
## Summary
(2-3 sentence overview)
## Key Findings
(bullet points of the most important facts)
## Details
(more depth on the most interesting points)
## Sources
(list URLs you found)

Be factual and concise. Only report what you found."""


# ════════════════════════════════════════════════════════════════════════════════
# RESEARCH TOOLS
# ════════════════════════════════════════════════════════════════════════════════

research_tools = [
    {
        "name": "web_search",
        "description": "Search the web for information on a topic. Use this first.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string"}
            },
            "required": ["query"]
        }
    },
    {
        "name": "fetch_page",
        "description": "Read the full content of a webpage URL for more detail.",
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {"type": "string"}
            },
            "required": ["url"]
        }
    }
]

def web_search(query: str) -> str:
    try:
        r    = requests.get(
            "https://api.duckduckgo.com/",
            params={"q": query, "format": "json", "no_html": 1, "skip_disambig": 1},
            timeout=10
        )
        data = r.json()
        results = []
        if data.get("AbstractText"):
            results.append(f"Summary: {data['AbstractText']}")
            results.append(f"Source: {data.get('AbstractURL', '')}")
        for topic in data.get("RelatedTopics", [])[:5]:
            if isinstance(topic, dict) and topic.get("Text"):
                results.append(f"- {topic['Text']}")
        return "\n".join(results) if results else f"No direct results for '{query}'."
    except Exception as e:
        return f"Search error: {str(e)}"

def fetch_page(url: str) -> str:
    try:
        r    = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        lines = [l for l in soup.get_text("\n", strip=True).splitlines() if len(l.strip()) > 40]
        return "\n".join(lines[:80]) or "No readable content found."
    except requests.exceptions.Timeout:
        return "Error: Page timed out."
    except requests.exceptions.HTTPError as e:
        return f"Error: HTTP {e.response.status_code}"
    except Exception as e:
        return f"Fetch error: {str(e)}"

def run_research_tool(name: str, inputs: dict) -> str:
    try:
        if name == "web_search": return web_search(inputs["query"])
        if name == "fetch_page": return fetch_page(inputs["url"])
        return f"Unknown tool: {name}"
    except Exception as e:
        return f"Tool '{name}' failed: {str(e)}"


# ════════════════════════════════════════════════════════════════════════════════
# SESSION STATE INITIALISATION
# ════════════════════════════════════════════════════════════════════════════════

# chat_docs: list of {"name": str, "text": str} — the active document pool
# chat_collection: ChromaDB collection built from chat_docs
# doc_messages: conversation history for Document Chat tab
# research_messages: conversation history for Research Agent tab
# research_results: list of {"label": str, "text": str} collected during research

if "chat_docs"          not in st.session_state:
    # Load default docs (Hassan's resume) on first run
    default_docs = load_default_docs("my_docs")
    st.session_state.chat_docs = default_docs

if "chat_collection"    not in st.session_state:
    if st.session_state.chat_docs:
        st.session_state.chat_collection = build_collection_from_docs(
            st.session_state.chat_docs, "chat_docs"
        )
    else:
        st.session_state.chat_collection = None

if "doc_messages"       not in st.session_state:
    st.session_state.doc_messages       = []

if "research_messages"  not in st.session_state:
    st.session_state.research_messages  = []

if "research_log"       not in st.session_state:
    st.session_state.research_log       = []

# research_results: accumulated text snippets from research
# each entry: {"label": str, "text": str, "added": bool}
if "research_results"   not in st.session_state:
    st.session_state.research_results   = []

# switch_to_chat: flag set when user clicks "Start Chatting" in Research tab
if "switch_to_chat"     not in st.session_state:
    st.session_state.switch_to_chat     = False

# force_doc_title: overrides title to "Document Chat" for one render cycle
# after Start Chatting is clicked, regardless of what query params say
if "force_doc_title"    not in st.session_state:
    st.session_state.force_doc_title    = False

# research_history: list of completed research runs
# each entry: {"question": str, "report": str, "steps": int, "results": list}
if "research_history"   not in st.session_state:
    st.session_state.research_history   = []

if "pending_doc_query"  not in st.session_state:
    st.session_state.pending_doc_query  = None   # shown at top while processing

if "pending_res_query"  not in st.session_state:
    st.session_state.pending_res_query  = None   # shown at top while processing


# ════════════════════════════════════════════════════════════════════════════════
# TABS  —  Document Chat is Tab 1 (default)
# Active tab is detected via query params set by a JS MutationObserver below.
# ════════════════════════════════════════════════════════════════════════════════

# Read active tab from query params (set by JS observer)
# Default to "doc" on first load
active_tab = st.query_params.get("tab", "doc")

tab1, tab2 = st.tabs(["📄 Document Chat", "🔍 Research Agent"])

# ── JS MutationObserver — watches tab clicks and syncs to query params ────────
# When the user clicks a tab, the selected tab's aria-selected attribute changes
# to "true". The observer detects this and updates the URL query param,
# which causes Streamlit to rerun and pick up the new active_tab value.
components.html(
    """
    <script>
    (function() {
        function syncTab() {
            const tabs = window.parent.document.querySelectorAll(
                'button[data-baseweb="tab"]'
            );
            if (!tabs.length) return;
            const activeIndex = Array.from(tabs).findIndex(
                t => t.getAttribute("aria-selected") === "true"
            );
            const paramValue = activeIndex === 1 ? "research" : "doc";
            const url = new URL(window.parent.location.href);
            if (url.searchParams.get("tab") !== paramValue) {
                url.searchParams.set("tab", paramValue);
                window.parent.history.replaceState({}, "", url.toString());
            }
        }
        // Observe the whole document for attribute changes on tab buttons
        const observer = new MutationObserver(syncTab);
        observer.observe(window.parent.document.body, {
            subtree: true, attributes: true, attributeFilter: ["aria-selected"]
        });
        syncTab(); // run once immediately
    })();
    </script>
    """,
    height=0
)


# ════════════════════════════════════════════════════════════════════════════════
# TAB 1 — DOCUMENT CHAT
# ════════════════════════════════════════════════════════════════════════════════

with tab1:

    # Title logic:
    # force_doc_title takes priority — set when Start Chatting is clicked.
    # It overrides the query param for one render cycle so the title is
    # correct before the MutationObserver has a chance to sync the URL.

    if st.session_state.force_doc_title or active_tab == "doc":
        st.title("📄 Document Chat")
        st.session_state.force_doc_title = False
    else:
        st.title("🧠 Learn Smarter")


    # ── JS tab switch when user clicks "Start Chatting" from Research tab ──────
    # Streamlit doesn't expose a tab-switching API, so we inject a tiny JS
    # snippet that clicks the first tab button (Document Chat) in the DOM.
    if st.session_state.switch_to_chat:
        st.session_state.switch_to_chat  = False
        st.session_state.force_doc_title = True   # ensures title shows "Document Chat"
        research_doc_names = [
            d["name"].replace("research:", "")
            for d in st.session_state.chat_docs
            if d["name"].startswith("research:")
        ]
        st.success(
            f"✅ Ready to chat! {len(research_doc_names)} research document(s) loaded: "
            + ", ".join(research_doc_names)
        )
        # Click the first tab (Document Chat) via JS
        components.html(
            """
            <script>
            const tabs = window.parent.document.querySelectorAll(
                'button[data-baseweb="tab"]'
            );
            if (tabs.length > 0) { tabs[0].click(); }
            </script>
            """,
            height=0
        )

    # ── Welcome message explaining the sample doc and options ─────────────────
    # NOTE: This info box is shown when the default sample document is the only
    # document loaded. It explains what the user can do next.
    doc_names = [d["name"] for d in st.session_state.chat_docs]

    # Auto-remove the default resume as soon as any other doc is present
    DEFAULT_NAMES = {"hassan_resume.txt", "hassan_resume.md"}
    has_non_default = any(n not in DEFAULT_NAMES for n in doc_names)
    if has_non_default:
        before = len(st.session_state.chat_docs)
        st.session_state.chat_docs = [
            d for d in st.session_state.chat_docs
            if d["name"] not in DEFAULT_NAMES
        ]
        if len(st.session_state.chat_docs) != before:
            doc_names = [d["name"] for d in st.session_state.chat_docs]
            rebuild_chat_collection()

    is_default_only = (
        len(st.session_state.chat_docs) > 0
        and all(
            name in DEFAULT_NAMES
            for name in doc_names
        )
    )

    # ── Dynamic subtitle based on active docs ─────────────────────────────────
    if doc_names:
        if is_default_only:
            st.caption("🧑‍💼 Chatting with: Hassan's resume (sample document)")
        else:
            clean_names = [
                d["name"].replace("research:", "🔍 ") if d["name"].startswith("research:")
                else d["name"]
                for d in st.session_state.chat_docs
            ]
            st.caption(f"💬 Chatting with: {', '.join(clean_names)}")

    # ── Sidebar ───────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown(
        '''<div style="display:flex;align-items:center;gap:12px;padding:8px 0 14px;">
          <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100" width="60" height="60" style="flex-shrink:0;">
            <rect width="100" height="100" rx="22" fill="#0f172a"/>
            <circle cx="50" cy="50" r="10" fill="#fbbf24"/>
            <circle cx="50" cy="50" r="6" fill="#fbbf24"/>
            <line x1="50" y1="40" x2="50" y2="18" stroke="#fbbf24" stroke-width="2.5" stroke-linecap="round"/>
            <line x1="50" y1="60" x2="50" y2="82" stroke="#fbbf24" stroke-width="2.5" stroke-linecap="round"/>
            <line x1="40" y1="50" x2="18" y2="50" stroke="#fbbf24" stroke-width="2.5" stroke-linecap="round"/>
            <line x1="60" y1="50" x2="82" y2="50" stroke="#fbbf24" stroke-width="2.5" stroke-linecap="round"/>
            <line x1="43" y1="43" x2="25" y2="25" stroke="#fbbf24" stroke-width="2" stroke-linecap="round"/>
            <line x1="57" y1="43" x2="75" y2="25" stroke="#fbbf24" stroke-width="2" stroke-linecap="round"/>
            <line x1="43" y1="57" x2="25" y2="75" stroke="#fbbf24" stroke-width="2" stroke-linecap="round"/>
            <line x1="57" y1="57" x2="75" y2="75" stroke="#fbbf24" stroke-width="2" stroke-linecap="round"/>
            <circle cx="50" cy="16" r="5.5" fill="#fbbf24" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="50" cy="84" r="5.5" fill="#fbbf24" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="16" cy="50" r="5.5" fill="#fbbf24" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="84" cy="50" r="5.5" fill="#fbbf24" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="25" cy="25" r="4" fill="#0f172a" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="75" cy="25" r="4" fill="#0f172a" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="25" cy="75" r="4" fill="#0f172a" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="75" cy="75" r="4" fill="#0f172a" stroke="#fbbf24" stroke-width="1.5"/>
          </svg>
          <div style="line-height:1.3;">
            <div style="font-size:1rem;font-weight:700;color:#e8f0fe;font-family:sans-serif;">Learn Smarter</div>
            <div style="font-size:0.65rem;color:#fbbf24 !important;letter-spacing:1.5px;font-family:sans-serif;font-weight:600;">AI-POWERED</div>
          </div>
        </div>''',
        unsafe_allow_html=True
    )
        st.divider()

        # ── Files currently in chat ───────────────────────────────────────────
        st.subheader("📂 Files in Chat")

        MAX_BYTES = 5 * 1024 * 1024  # 5 MB
        used_bytes = sum(len(d["text"].encode("utf-8")) for d in st.session_state.chat_docs)
        used_mb    = used_bytes / (1024 * 1024)
        pct        = min(used_bytes / MAX_BYTES, 1.0)

        # Storage meter
        st.progress(pct, text=f"Storage: {used_mb:.2f} MB / 5.00 MB")
        if pct >= 1.0:
            st.error("⚠️ **Storage full.** Remove a document to add more.")
        elif pct >= 0.8:
            st.warning(f"⚠️ Nearly full — {(1-pct)*100:.0f}% remaining.")

        if st.session_state.chat_docs:
            for i, doc in enumerate(st.session_state.chat_docs):
                is_research = doc["name"].startswith("research:")
                is_default  = doc["name"] in ["hassan_resume.txt", "hassan_resume.md"]
                if is_research:
                    display_name = doc["name"].replace("research:", "").strip()
                    icon = "🔍"
                elif is_default:
                    display_name = doc["name"]
                    icon = "🧑‍💼"
                else:
                    ext  = os.path.splitext(doc["name"])[1].lower()
                    icon = {".pdf": "📕", ".docx": "📘", ".md": "📝"}.get(ext, "📄")
                    display_name = doc["name"]
                doc_kb = len(doc["text"].encode("utf-8")) / 1024
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.caption(f"{icon} {display_name[:28]}{'…' if len(display_name) > 28 else ''} · {doc_kb:.0f} KB")
                with col2:
                    if st.button("✕", key=f"remove_{i}", help=f"Remove {display_name}"):
                        st.session_state.chat_docs.pop(i)
                        rebuild_chat_collection()
                        st.rerun()
        else:
            st.caption("No documents loaded.")

        st.divider()

        # ── Upload your own document ──────────────────────────────────────────
        st.subheader("Upload Documents")

        uploaded_file = st.file_uploader(
            "Choose files",
            type=["txt", "md", "pdf", "docx"],
            help="Supported: .txt  .md  .pdf  .docx"
        )

        st.info(
            "**Accepted formats:**\n"
            "- `.txt` — plain text\n"
            "- `.md` — markdown\n"
            "- `.pdf` — PDF (must have selectable text, not scanned)\n"
            "- `.docx` — Word document\n\n"
            "Uploaded files and research results are **combined** in chat."
        )

        if uploaded_file:
            already_loaded = any(
                d["name"] == uploaded_file.name
                and not d["name"].startswith("research:")
                for d in st.session_state.chat_docs
            )
            # Count non-default docs and check combined storage
            MAX_BYTES = 5 * 1024 * 1024
            used_bytes = sum(len(d["text"].encode("utf-8")) for d in st.session_state.chat_docs)
            non_default_count = len([
                d for d in st.session_state.chat_docs
                if d["name"] not in ["hassan_resume.txt", "hassan_resume.md"]
            ])
            if already_loaded:
                st.caption(f"✅ {uploaded_file.name} already in chat.")
            elif non_default_count >= 3:
                st.warning(
                    "⚠️ **Document limit reached.** This free app supports a maximum of "
                    "**3 documents** in chat at a time. Remove a document from "
                    "Files in Chat before uploading another."
                )
            else:
                with st.spinner("Indexing..."):
                    text = extract_text_from_file(uploaded_file)
                if text.strip():
                    new_bytes = len(text.encode("utf-8"))
                    if used_bytes + new_bytes > MAX_BYTES:
                        st.warning(
                            f"⚠️ **Storage limit reached.** This free app allows a combined "
                            f"maximum of **5 MB** across all documents. "
                            f"You are using **{used_bytes/1024/1024:.2f} MB** — this file would "
                            f"add **{new_bytes/1024/1024:.2f} MB** and exceed the limit. "
                            f"Remove a document first."
                        )
                    else:
                        st.session_state.chat_docs = [
                            d for d in st.session_state.chat_docs
                            if d["name"] not in ["hassan_resume.txt", "hassan_resume.md"]
                        ]
                        st.session_state.chat_docs.append({
                            "name": uploaded_file.name,
                            "text": text
                        })
                        rebuild_chat_collection()
                        st.rerun()
                else:
                    st.error(
                        "Could not extract text. Check that your file:\n"
                        "- Is not password protected\n"
                        "- Contains selectable text (not a scanned image)\n"
                        "- Is not corrupted"
                    )


        st.divider()

        # ── Controls ──────────────────────────────────────────────────────────
        if st.button("🗑️ Clear conversation", use_container_width=True, key="clear_doc"):
            st.session_state.doc_messages = []
            st.rerun()

        if st.button("🔄 Reset to default document", use_container_width=True):
            # Remove ALL uploaded files, research docs and links, then load default
            default_docs = load_default_docs("my_docs")
            st.session_state.chat_docs       = default_docs
            st.session_state.doc_messages    = []
            st.session_state.research_results = []
            rebuild_chat_collection()
            st.rerun()

        st.divider()
        show_context = st.toggle("Show retrieved context", value=False)
        st.caption("See which document sections were used for each answer.")

    # ── Chat input — TOP of chat area ────────────────────────────────────────
    chat_disabled    = st.session_state.chat_collection is None
    chat_placeholder = (
        "Ask a question about your documents..."
        if not chat_disabled
        else "Upload a document to start chatting..."
    )

    if prompt := st.chat_input(chat_placeholder, key="doc_input", disabled=chat_disabled):
        st.session_state.pending_doc_query = prompt
        st.rerun()

    # ── Show pending query at the top immediately while processing ────────────
    if st.session_state.pending_doc_query:
        pending = st.session_state.pending_doc_query
        with st.chat_message("user"):
            st.markdown(f"**{pending}**")
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                context = retrieve_context(st.session_state.chat_collection, pending)
                doc_label = ", ".join(doc_names) if doc_names else "the documents"
                augmented = f"""Here is relevant information from {doc_label}:

{context}

---
Question: {pending}"""
                history = []
                for msg in [m for m in st.session_state.doc_messages if m["role"] != "_context"]:
                    history.append({"role": msg["role"], "content": msg["content"]})
                history.append({"role": "user", "content": augmented})
                active_system_prompt = build_chat_system_prompt(doc_names)
                full_response = ""
                with client.messages.stream(
                    model="claude-sonnet-4-6",
                    max_tokens=1024,
                    system=active_system_prompt,
                    messages=history
                ) as stream:
                    for chunk in stream.text_stream:
                        full_response += chunk
                st.markdown(full_response)

        st.session_state.doc_messages.append({"role": "user", "content": pending})
        st.session_state.doc_messages.append({"role": "assistant", "content": full_response})
        if show_context:
            st.session_state.doc_messages.append({"role": "_context", "content": context})
        st.session_state.pending_doc_query = None
        st.rerun()

    # ── Build grouped pairs: [(user_msg, assistant_msg), ...] ──────────────
    msgs      = st.session_state.doc_messages
    real_msgs = [m for m in msgs if m["role"] != "_context"]

    # Group into (user, assistant) pairs
    pairs = []
    i = 0
    while i < len(real_msgs):
        if real_msgs[i]["role"] == "user":
            user_msg = real_msgs[i]
            asst_msg = real_msgs[i + 1] if i + 1 < len(real_msgs) and real_msgs[i + 1]["role"] == "assistant" else None
            pairs.append((user_msg, asst_msg))
            i += 2 if asst_msg else 1
        else:
            i += 1

    if pairs:
        # ── Latest pair just below the input ────────────────────────────────
        latest_user, latest_asst = pairs[-1]
        st.divider()
        with st.chat_message("user"):
            st.markdown(latest_user["content"])
        if latest_asst:
            with st.chat_message("assistant"):
                st.markdown(latest_asst["content"])

        if show_context:
            ctx = next((m["content"] for m in msgs if m["role"] == "_context"), None)
            if ctx:
                with st.expander("📄 Retrieved document sections", expanded=False):
                    st.text(ctx)
                chunk_count = len(ctx.split("[Document")) - 1
                if chunk_count > 0:
                    st.caption(f"Retrieved {chunk_count} document sections")

        # ── Historic pairs below, newest first ───────────────────────────────
        historic_pairs = pairs[:-1]
        if historic_pairs:
            st.divider()
            st.caption("🕘 **Previous exchanges**")
            for user_msg, asst_msg in reversed(historic_pairs):
                with st.container(border=True):
                    with st.chat_message("user"):
                        st.markdown(user_msg["content"])
                    if asst_msg:
                        with st.chat_message("assistant"):
                            st.markdown(asst_msg["content"])

    # ── Welcome / empty state ─────────────────────────────────────────────────
    elif is_default_only:
        with st.chat_message("assistant"):
            st.markdown(
                "👋 **Welcome to Learn Smarter!**\n\n"
                "I've loaded **Hassan's resume** as a sample. Chat with it now, "
                "upload your own document from the sidebar, or switch to the "
                "**Research Agent** tab to research any topic."
            )
    elif not st.session_state.chat_docs:
        with st.chat_message("assistant"):
            st.markdown(
                "No documents loaded yet. Upload a document from the sidebar "
                "or add research results from the Research Agent tab."
            )


# ════════════════════════════════════════════════════════════════════════════════
# TAB 2 — RESEARCH AGENT
# ════════════════════════════════════════════════════════════════════════════════

with tab2:
    st.title("🔍 Research Agent")
    st.caption(
        "Research any topic. The agent searches the web and compiles a report. "
        "Add results to Document Chat to ask follow-up questions."
    )
    st.info("⚠️ Only the **2 most recent** research sessions are kept. Older ones are removed automatically.")

    st.divider()

    # ── Prompt input ──────────────────────────────────────────────────────────
    research_prompt = st.chat_input("Enter a research question...", key="research_input")

    # ── Show pending question + spinner at TOP — history remains visible below ──
    if st.session_state.pending_res_query:
        with st.chat_message("user"):
            st.markdown(f"**{st.session_state.pending_res_query}**")
        with st.chat_message("assistant"):
            st.info("🧠 Research in progress — searching the web...")
        st.divider()

    # ── Action buttons — above the history ───────────────────────────────────
    if st.session_state.research_history:
        col1, col2 = st.columns([3, 1])
        with col1:
            research_added = [
                d["name"] for d in st.session_state.chat_docs
                if d["name"].startswith("research:")
            ]
            if research_added:
                if st.button("💬 Start Chatting →", type="primary", use_container_width=True, key="start_chatting"):
                    st.session_state.switch_to_chat = True
                    st.rerun()
        with col2:
            if st.button("🗑️ Clear history", use_container_width=True, key="clear_history"):
                st.session_state.research_history  = []
                st.session_state.research_results  = []
                st.session_state.research_messages = []
                st.session_state.research_log      = []
                st.rerun()

        st.divider()

    # ── Show history: latest on top, question + response grouped ─────────────
    if st.session_state.research_history:
        MAX_DOCS = 3
        for h_idx, entry in enumerate(reversed(st.session_state.research_history)):
            real_idx  = len(st.session_state.research_history) - 1 - h_idx

            with st.container(border=True):
                # Question on top
                with st.chat_message("user"):
                    st.markdown(f"**{entry['question']}**")

                # Response below question
                with st.chat_message("assistant"):
                    st.markdown(entry["report"])

                # Sources / result links below response
                link_results = [r for r in entry["results"] if r.get("type") != "report"]
                if link_results:
                    st.caption("**Sources — add to Document Chat:**")
                    # Count non-default docs currently in chat
                    non_default_count = len([
                        d for d in st.session_state.chat_docs
                        if d["name"] not in ["hassan_resume.txt", "hassan_resume.md"]
                    ])
                    for r_idx, result in enumerate(link_results):
                        rtype  = result.get("type", "url")
                        icon   = "🔗" if rtype == "url" else "🔎"
                        label  = result["label"]
                        already = any(
                            d["name"] == f"research:{label}"
                            for d in st.session_state.chat_docs
                        )
                        col1, col2 = st.columns([5, 1])
                        with col1:
                            st.caption(f"{'✅' if already else icon} {label[:72]}")
                        with col2:
                            if already:
                                st.caption("Added ✓")
                            elif non_default_count >= MAX_DOCS:
                                st.caption("⚠️ Limit")
                            else:
                                # Also check 5MB combined storage limit
                                MAX_BYTES    = 5 * 1024 * 1024
                                used_bytes   = sum(len(d["text"].encode("utf-8")) for d in st.session_state.chat_docs)
                                new_bytes    = len(result["text"].encode("utf-8"))
                                over_storage = used_bytes + new_bytes > MAX_BYTES
                                if over_storage:
                                    st.caption("⚠️ Full")
                                elif st.button(
                                    "Add",
                                    key=f"hist_{real_idx}_{r_idx}",
                                    use_container_width=True
                                ):
                                    st.session_state.chat_docs = [
                                        d for d in st.session_state.chat_docs
                                        if d["name"] not in ["hassan_resume.txt", "hassan_resume.md"]
                                    ]
                                    st.session_state.chat_docs.append({
                                        "name": f"research:{label}",
                                        "text": result["text"]
                                    })
                                    rebuild_chat_collection()
                                    st.rerun()

                    # Show limit warnings below sources
                    used_bytes_now = sum(len(d["text"].encode("utf-8")) for d in st.session_state.chat_docs)
                    if non_default_count >= MAX_DOCS:
                        st.warning(
                            "⚠️ **Document limit reached.** This free app supports a maximum of "
                            "**3 documents** in chat at a time. Remove a document from the "
                            "sidebar to add another."
                        )
                    elif used_bytes_now >= 5 * 1024 * 1024:
                        st.warning(
                            "⚠️ **Storage limit reached (5 MB).** This free app limits combined "
                            "document storage to 5 MB. Remove a document from the sidebar to add more."
                        )

    # ── Run research when prompt submitted ────────────────────────────────────
    if research_prompt:
        st.session_state.pending_res_query = research_prompt
        st.rerun()

    if st.session_state.pending_res_query:
        research_prompt = st.session_state.pending_res_query
        # Question already shown at top — now run the agent
        st.session_state.research_messages.append({"role": "user", "content": research_prompt})
        st.session_state.research_log     = []
        st.session_state.research_results = []
        st.session_state.pending_res_query = None

        messages         = [{"role": "user", "content": research_prompt}]
        step             = 1
        status_container = st.empty()

        # Track URLs and text fetched during this research run
        fetched_pages  = {}  # url → text
        search_results = {}  # query → text

        status_container.info("🧠 Research in progress — gathering sources...")

        with st.chat_message("assistant"):
            response_placeholder = st.empty()

            while True:
                # Update status while Claude is thinking between steps
                if step == 1:
                    status_container.info("🧠 Research in progress — analysing your question...")
                else:
                    status_container.info(f"🧠 Research in progress — processing step {step}...")

                response = client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=4096,
                    system=RESEARCH_SYSTEM_PROMPT,
                    tools=research_tools,
                    messages=messages
                )

                if response.stop_reason == "end_turn":
                    final_text = ""
                    for block in response.content:
                        if hasattr(block, "text"):
                            final_text = block.text
                            break

                    status_container.success(f"✅ Research complete — {step - 1} steps taken")
                    response_placeholder.markdown(final_text)
                    st.session_state.research_messages.append({
                        "role": "assistant", "content": final_text
                    })

                    # ── Build result links for this session ───────────────────
                    run_results = []
                    # Add each fetched page as a result link
                    for url, text in fetched_pages.items():
                        entry = {"label": url, "type": "url", "text": text}
                        run_results.append(entry)
                    # Add search queries as result links
                    for query, text in search_results.items():
                        entry = {"label": f"Search: {query}", "type": "search", "text": text}
                        run_results.append(entry)
                    # Add the full report as a result item
                    run_results.append({
                        "label": f"Full report",
                        "type":  "report",
                        "text":  final_text
                    })

                    # ── Save to research history — keep only 2 most recent ─────
                    st.session_state.research_history.append({
                        "question": research_prompt,
                        "report":   final_text,
                        "steps":    step - 1,
                        "results":  run_results
                    })
                    # Trim to last 2 sessions
                    if len(st.session_state.research_history) > 2:
                        st.session_state.research_history = st.session_state.research_history[-2:]
                    st.session_state.research_results = run_results
                    break

                if response.stop_reason == "tool_use":
                    messages.append({"role": "assistant", "content": response.content})
                    tool_results = []

                    for block in response.content:
                        if block.type == "tool_use":
                            label = (
                                f"🔎 Searching: *{block.input.get('query', '')}*"
                                if block.name == "web_search"
                                else f"📖 Reading: *{block.input.get('url', '')[:60]}...*"
                            )
                            status_container.info(f"Step {step}: {label}")
                            log_entry = f"Step {step} — {block.name}: {list(block.input.values())[0][:80]}"
                            st.session_state.research_log.append(log_entry)

                            result = run_research_tool(block.name, block.input)

                            # Store fetched pages and search results for history
                            if block.name == "fetch_page":
                                fetched_pages[block.input.get("url", f"page_{step}")] = result
                            elif block.name == "web_search":
                                query = block.input.get("query", f"search_{step}")
                                if query not in search_results:
                                    search_results[query] = result

                            tool_results.append({
                                "type": "tool_result",
                                "tool_use_id": block.id,
                                "content": result
                            })
                            step += 1

                    messages.append({"role": "user", "content": tool_results})

                else:
                    status_container.empty()
                    break

        st.caption(f"Research completed in {step - 1} steps")
        st.rerun()  # refresh to show the Add to Chat buttons immediately
