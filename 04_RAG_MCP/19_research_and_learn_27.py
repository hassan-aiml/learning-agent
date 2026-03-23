import os
import io
import requests
import chromadb
import streamlit as st
import streamlit.components.v1 as components
from dotenv import load_dotenv
from anthropic import Anthropic
from bs4 import BeautifulSoup
from sentence_transformers import SentenceTransformer

# Load API key — works both locally (.env file) and on Streamlit Cloud (st.secrets)
load_dotenv()
if "ANTHROPIC_API_KEY" not in os.environ:
    os.environ["ANTHROPIC_API_KEY"] = st.secrets.get("ANTHROPIC_API_KEY", "")

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Learn Smart",
    page_icon="🧠",
    layout="centered"
)

# ── Global styles: tab colors + mobile viewport ───────────────────────────────
# st.html() is used here because st.markdown() can strip or misrender
# <style> blocks in some Streamlit versions, causing CSS to appear as text.
st.html("""
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>

/* ── Learn Smart — Blue Light theme ── */
:root {
    --navy:          #1e3a5f;
    --navy-mid:      #162d4a;
    --blue:          #1d4ed8;
    --blue-mid:      #1e40af;
    --blue-light:    #eff6ff;
    --blue-pale:     #f0f7ff;
    --blue-glow:     rgba(29,78,216,0.12);
    --sky:           #60a5fa;
    --sky-light:     #bfdbfe;
    --card-white:    #ffffff;
    --border:        #dbeafe;
    --text-primary:  #1e3a5f;
    --text-secondary:#3b5f8a;
    --text-hint:     #94a3b8;
    --text-sidebar:  #e8f0fe;
}

/* ── Page background: soft ice blue ── */
html, body,
[data-testid="stAppViewContainer"],
[data-testid="stAppViewBlockContainer"],
.stApp, .main {
    background-color: var(--blue-pale) !important;
    color: var(--text-primary) !important;
}

/* ── Content card: white ── */
.block-container {
    background: var(--card-white) !important;
    padding: 1.5rem 2rem !important;
    box-shadow: 0 1px 4px rgba(29,78,216,0.08), 0 2px 10px rgba(29,78,216,0.05) !important;
}
@media (min-width: 641px) {
    .block-container {
        border-radius: 12px !important;
        margin-top: 1rem !important;
        border: 1px solid var(--border) !important;
    }
}

/* ── Body text ── */
.block-container p,
.block-container li,
.block-container span,
.block-container label,
[data-testid="stMarkdownContainer"] p {
    color: var(--text-primary) !important;
}

/* ── Page titles ── */
h1, .block-container h1 {
    color: var(--navy) !important;
    font-weight: 700 !important;
    border-bottom: 2px solid var(--sky-light) !important;
    padding-bottom: 0.5rem !important;
    margin-bottom: 0.75rem !important;
}

/* ── Subheaders ── */
h2, .block-container h2 {
    color: var(--text-primary) !important;
    font-size: 1.1rem !important;
    font-weight: 600 !important;
}
h3, .block-container h3 {
    color: var(--blue) !important;
    font-size: 1rem !important;
    font-weight: 500 !important;
}

/* ── Captions ── */
.stCaption p, small {
    color: var(--text-secondary) !important;
}

/* ── Sidebar: deep navy ── */
section[data-testid="stSidebar"],
section[data-testid="stSidebar"] > div,
section[data-testid="stSidebar"] > div > div {
    background: var(--navy) !important;
    background-color: var(--navy) !important;
    border-right: 1px solid var(--navy-mid) !important;
}
section[data-testid="stSidebar"] * {
    color: var(--text-sidebar) !important;
}
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: var(--sky) !important;
}
section[data-testid="stSidebar"] .stCaption p {
    color: #93c5fd !important;
}
section[data-testid="stSidebar"] hr {
    border-color: #2a4a70 !important;
}
section[data-testid="stSidebar"] [data-testid="stAlert"] {
    background: var(--navy-mid) !important;
    border-left: 3px solid var(--sky) !important;
}
section[data-testid="stSidebar"] [data-testid="stMetricValue"] {
    color: var(--sky) !important;
    font-size: 1.4rem !important;
}

/* ── All buttons: theme blue, white text ── */
/* Default state: solid blue matching the theme accent */
.stButton > button,
[data-testid^="stBaseButton"],
button[data-testid^="stBaseButton"] {
    background-color: #1d4ed8 !important;
    background: #1d4ed8 !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    font-size: 0.875rem !important;
    box-shadow: 0 1px 3px rgba(29,78,216,0.3) !important;
    transition: all 0.2s !important;
    cursor: pointer !important;
}
/* Hover state: lighter sky blue — clear contrast with white text */
.stButton > button:hover,
[data-testid^="stBaseButton"]:hover,
button[data-testid^="stBaseButton"]:hover {
    background-color: #2563eb !important;
    background: #2563eb !important;
    color: #ffffff !important;
    box-shadow: 0 2px 6px rgba(29,78,216,0.4) !important;
    transform: translateY(-1px) !important;
}

/* ── Remove (✕) button: red on hover — clearly destructive ── */
.stButton > button[title*="Remove"]:hover,
.stButton > button[title*="remove"]:hover {
    background-color: #dc2626 !important;
    background: #dc2626 !important;
    color: #ffffff !important;
    box-shadow: 0 2px 6px rgba(220,38,38,0.4) !important;
}

/* ── Force white text on ALL button states everywhere ── */
.stButton > button *,
.stButton > button p,
.stButton > button span,
.stButton > button div,
[data-testid^="stBaseButton"] *,
[data-testid^="stBaseButton"] p,
[data-testid^="stBaseButton"] span {
    color: #ffffff !important;
}

/* ── Chat bubbles ── */
[data-testid="stChatMessageContent"] {
    border-radius: 10px !important;
    background: #f8fafc !important;
    border: 1px solid var(--border) !important;
    color: var(--text-primary) !important;
}
[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) [data-testid="stChatMessageContent"] {
    background: var(--blue-light) !important;
    border: 1px solid var(--sky-light) !important;
    border-left: 3px solid var(--blue) !important;
}

/* ── Chat avatar icons ── */
[data-testid="stChatMessageAvatarUser"] {
    background-color: #1d4ed8 !important;
    color: #ffffff !important;
}
[data-testid="stChatMessageAvatarAssistant"] {
    background-color: #e8edf2 !important;
    color: #1e3a5f !important;
}

/* ── Chat input ── */
[data-testid="stChatInput"],
[data-testid="stChatInput"] textarea {
    background: var(--card-white) !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 24px !important;
    color: var(--text-primary) !important;
}
[data-testid="stChatInput"]:focus-within {
    border-color: var(--blue) !important;
    box-shadow: 0 0 0 2px var(--blue-glow) !important;
    outline: none !important;
}

/* ── Remove purple/default focus ring on ALL inputs ── */
*:focus, *:focus-visible {
    outline-color: var(--blue) !important;
    outline-width: 2px !important;
}
textarea:focus, input:focus {
    outline: none !important;
    box-shadow: 0 0 0 2px var(--blue-glow) !important;
    border-color: var(--blue) !important;
}

/* ── File uploader drag and drop area ── */
[data-testid="stFileUploader"] {
    background: var(--navy) !important;
    border-radius: 8px !important;
}
[data-testid="stFileUploaderDropzone"] {
    background: var(--navy) !important;
    border: 2px dashed var(--sky) !important;
    border-radius: 8px !important;
}
[data-testid="stFileUploaderDropzone"] * {
    color: #e8f0fe !important;
}
[data-testid="stFileUploaderDropzone"] p,
[data-testid="stFileUploaderDropzone"] span,
[data-testid="stFileUploaderDropzone"] small {
    color: #bfdbfe !important;
}
[data-testid="stFileUploaderDropzone"]:hover {
    border-color: var(--blue) !important;
    background: var(--navy-mid) !important;
}

/* ── Alerts ── */
[data-testid="stAlert"] {
    border-radius: 8px !important;
    background: var(--blue-light) !important;
    border: 1px solid var(--border) !important;
}

/* ── Expanders ── */
[data-testid="stExpander"] {
    background: var(--card-white) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
}
[data-testid="stExpander"] summary {
    color: var(--blue) !important;
    font-weight: 500 !important;
}

/* ── Dividers ── */
hr { border-color: var(--border) !important; }

/* ── Metrics ── */
[data-testid="stMetricValue"] { color: var(--blue) !important; }
[data-testid="stMetricLabel"] { color: var(--text-secondary) !important; }

/* ── Tab bar ── */
[data-baseweb="tab-list"] {
    background: var(--card-white) !important;
    border-bottom: 2px solid var(--border) !important;
    padding: 0 !important;
}
button[data-baseweb="tab"] {
    color: var(--text-hint) !important;
    font-weight: 500 !important;
}
button[data-baseweb="tab"]:nth-of-type(1):hover {
    color: var(--blue) !important;
    background: var(--blue-light) !important;
    border-radius: 6px 6px 0 0 !important;
}
button[data-baseweb="tab"]:nth-of-type(1)[aria-selected="true"] {
    background: var(--blue-light) !important;
    border-bottom: 3px solid var(--blue) !important;
    color: var(--blue) !important;
    font-weight: 700 !important;
    border-radius: 6px 6px 0 0 !important;
}
button[data-baseweb="tab"]:nth-of-type(2):hover {
    color: #0891b2 !important;
    background: #ecfeff !important;
    border-radius: 6px 6px 0 0 !important;
}
button[data-baseweb="tab"]:nth-of-type(2)[aria-selected="true"] {
    background: #ecfeff !important;
    border-bottom: 3px solid #0891b2 !important;
    color: #0891b2 !important;
    font-weight: 700 !important;
    border-radius: 6px 6px 0 0 !important;
}

/* ── Toggle ── */
[data-testid="stToggle"] input:checked + div {
    background: var(--blue) !important;
}

/* ── Alert accents ── */
[data-testid="stAlert"][kind="success"] { border-left: 3px solid #059669 !important; }
[data-testid="stAlert"][kind="warning"]  { border-left: 3px solid #d97706 !important; }

/* ── Mobile ── */
@media (max-width: 640px) {
    .block-container { padding: 0.75rem 0.6rem !important; border-radius: 0 !important; border: none !important; margin-top: 0 !important; }
    [data-testid="stTabs"] { margin-top: 0 !important; padding-top: 0 !important; }
    [data-baseweb="tab-list"] { position: sticky !important; top: 0 !important; z-index: 100 !important; width: 100% !important; margin: 0 !important; padding: 0 !important; display: flex !important; }
    button[data-baseweb="tab"] { flex: 1 !important; min-height: 48px !important; font-size: 0.8rem !important; padding: 0 4px !important; justify-content: center !important; touch-action: manipulation !important; }
    h1 { font-size: 1.2rem !important; margin-top: 0.5rem !important; }
    .stButton > button { font-size: 0.85rem !important; }
    section[data-testid="stSidebar"] { min-width: 260px !important; }
}

</style>
""")


# ── Cached resources ──────────────────────────────────────────────────────────
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

@st.cache_resource
def load_claude_client():
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        st.error("ANTHROPIC_API_KEY not found. Add it to your .env file locally or to Streamlit secrets in the cloud.")
        st.stop()
    return Anthropic(api_key=api_key)

model  = load_embedding_model()
client = load_claude_client()


# ════════════════════════════════════════════════════════════════════════════════
# SHARED UTILITIES
# ════════════════════════════════════════════════════════════════════════════════

def chunk_text(text: str, chunk_size: int = 300, overlap: int = 50) -> list:
    """Split text into overlapping chunks for RAG retrieval."""
    words  = text.split()
    chunks = []
    start  = 0
    while start < len(words):
        chunk = " ".join(words[start:start + chunk_size])
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def extract_text_from_file(uploaded_file) -> str:
    """
    Extract plain text from an uploaded file.
    Supports: .txt, .md, .pdf, .docx
    """
    filename = uploaded_file.name.lower()

    if filename.endswith((".txt", ".md")):
        return uploaded_file.read().decode("utf-8")

    elif filename.endswith(".pdf"):
        try:
            import fitz
            pdf_bytes = uploaded_file.read()
            doc       = fitz.open(stream=pdf_bytes, filetype="pdf")
            return "\n\n".join([page.get_text() for page in doc])
        except ImportError:
            st.error("PyMuPDF not installed. Run: pip install pymupdf")
            return ""
        except Exception as e:
            st.error(f"Could not read PDF: {e}")
            return ""

    elif filename.endswith(".docx"):
        try:
            from docx import Document
            doc  = Document(io.BytesIO(uploaded_file.read()))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            st.error("python-docx not installed. Run: pip install python-docx")
            return ""
        except Exception as e:
            st.error(f"Could not read Word document: {e}")
            return ""

    else:
        st.error(f"Unsupported file type: {filename}")
        return ""


def build_collection_from_docs(doc_list: list, collection_name: str):
    """
    Build a ChromaDB collection from a list of
    {"name": str, "text": str} dicts.
    Returns the collection or None if empty.
    """
    all_ids   = []
    all_texts = []

    for doc in doc_list:
        chunks = chunk_text(doc["text"])
        for i, chunk in enumerate(chunks):
            all_ids.append(f"{doc['name']}_chunk_{i}")
            all_texts.append(chunk)

    if not all_texts:
        return None

    base_dir = os.path.dirname(os.path.abspath(__file__))
    chroma_client = chromadb.PersistentClient(path=os.path.join(base_dir, "knowledge_base"))
    try:
        chroma_client.delete_collection(collection_name)
    except:
        pass

    collection     = chroma_client.create_collection(collection_name)
    all_embeddings = model.encode(all_texts, show_progress_bar=False).tolist()
    collection.add(ids=all_ids, embeddings=all_embeddings, documents=all_texts)
    return collection


def retrieve_context(collection, query: str, top_k: int = 4) -> str:
    """Semantic search — returns formatted top-k chunks."""
    query_embedding = model.encode([query]).tolist()
    results = collection.query(query_embeddings=query_embedding, n_results=top_k)
    chunks  = results["documents"][0]
    return "\n\n".join([f"[Document {i+1}]: {chunk}" for i, chunk in enumerate(chunks)])


def rebuild_chat_collection():
    """
    Rebuild the active chat collection from all docs
    currently in st.session_state.chat_docs.
    Called whenever a doc is added or removed.
    """
    if not st.session_state.chat_docs:
        st.session_state.chat_collection = None
        return

    collection = build_collection_from_docs(
        st.session_state.chat_docs, "chat_docs"
    )
    st.session_state.chat_collection = collection


# ════════════════════════════════════════════════════════════════════════════════
# DEFAULT DOCUMENT — Hassan's resume from my_docs/
# ════════════════════════════════════════════════════════════════════════════════

def load_default_docs(folder: str = "my_docs") -> list:
    """
    Load all supported files from my_docs/ as a list of
    {"name": str, "text": str} dicts.
    # NOTE: The default document (Hassan's resume) is a sample to demonstrate
    # the Document Chat feature. Users can upload their own documents,
    # or use the Research Agent to find content and add it to the chat.
    """
    # Build path relative to this app file — works on Streamlit Cloud
    # where the working directory may differ from the app file location
    base_dir = os.path.dirname(os.path.abspath(__file__))
    folder   = os.path.join(base_dir, folder)

    if not os.path.exists(folder):
        return []

    docs = []
    for filename in os.listdir(folder):
        if not filename.endswith((".txt", ".md", ".pdf", ".docx")):
            continue
        filepath = os.path.join(folder, filename)
        ext      = filename.lower()
        try:
            if ext.endswith((".txt", ".md")):
                with open(filepath, "r", encoding="utf-8") as f:
                    text = f.read().strip()
            elif ext.endswith(".pdf"):
                import fitz
                doc  = fitz.open(filepath)
                text = "\n\n".join([page.get_text() for page in doc])
            elif ext.endswith(".docx"):
                from docx import Document
                doc  = Document(filepath)
                text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                continue
            if text.strip():
                docs.append({"name": filename, "text": text})
        except Exception:
            continue
    return docs


# ════════════════════════════════════════════════════════════════════════════════
# SYSTEM PROMPTS
# ════════════════════════════════════════════════════════════════════════════════

def build_chat_system_prompt(doc_names: list) -> str:
    """Dynamically build a system prompt based on active documents."""
    if not doc_names:
        return "You are a helpful assistant."

    # Special persona when only Hassan's resume is loaded
    if doc_names == ["hassan_resume.txt"] or doc_names == ["hassan_resume.md"]:
        return """You are a professional HR assistant helping recruiters and hiring
managers learn about Hassan M. Hai from his resume.
- Only answer using the resume context provided — never invent experience or skills
- If the resume doesn't contain enough information to answer, say so clearly
- Refer to the candidate by first name: Hassan
- Be objective, professional, and concise
- Do not speculate about things not mentioned in the resume"""

    names_str = ", ".join(doc_names)
    return f"""You are a helpful assistant answering questions based on the following
documents: {names_str}.

IMPORTANT RULES:
- Only answer using the document context provided — never invent information
- If the documents don't contain enough information, say so clearly
- Be concise, accurate, and helpful
- When referencing specific facts, mention which document they come from
- Do not speculate about things not mentioned in the documents"""


RESEARCH_SYSTEM_PROMPT = """You are a professional research assistant. Research topics
thoroughly and produce clear, structured reports.

Your process:
1. Search for the topic to get an overview
2. Fetch 1-2 useful URLs for more detail
3. Synthesise everything into a structured report

Report format:
# [Topic Title]
## Summary
(2-3 sentence overview)
## Key Findings
(bullet points of the most important facts)
## Details
(more depth on the most interesting points)
## Sources
(list URLs you found)

Be factual and concise. Only report what you found."""


# ════════════════════════════════════════════════════════════════════════════════
# RESEARCH TOOLS
# ════════════════════════════════════════════════════════════════════════════════

research_tools = [
    {
        "name": "web_search",
        "description": "Search the web for information on a topic. Use this first.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string"}
            },
            "required": ["query"]
        }
    },
    {
        "name": "fetch_page",
        "description": "Read the full content of a webpage URL for more detail.",
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {"type": "string"}
            },
            "required": ["url"]
        }
    }
]

def web_search(query: str) -> str:
    try:
        r    = requests.get(
            "https://api.duckduckgo.com/",
            params={"q": query, "format": "json", "no_html": 1, "skip_disambig": 1},
            timeout=10
        )
        data = r.json()
        results = []
        if data.get("AbstractText"):
            results.append(f"Summary: {data['AbstractText']}")
            results.append(f"Source: {data.get('AbstractURL', '')}")
        for topic in data.get("RelatedTopics", [])[:5]:
            if isinstance(topic, dict) and topic.get("Text"):
                results.append(f"- {topic['Text']}")
        return "\n".join(results) if results else f"No direct results for '{query}'."
    except Exception as e:
        return f"Search error: {str(e)}"

def fetch_page(url: str) -> str:
    try:
        r    = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        lines = [l for l in soup.get_text("\n", strip=True).splitlines() if len(l.strip()) > 40]
        return "\n".join(lines[:80]) or "No readable content found."
    except requests.exceptions.Timeout:
        return "Error: Page timed out."
    except requests.exceptions.HTTPError as e:
        return f"Error: HTTP {e.response.status_code}"
    except Exception as e:
        return f"Fetch error: {str(e)}"

def run_research_tool(name: str, inputs: dict) -> str:
    try:
        if name == "web_search": return web_search(inputs["query"])
        if name == "fetch_page": return fetch_page(inputs["url"])
        return f"Unknown tool: {name}"
    except Exception as e:
        return f"Tool '{name}' failed: {str(e)}"


# ════════════════════════════════════════════════════════════════════════════════
# SESSION STATE INITIALISATION
# ════════════════════════════════════════════════════════════════════════════════

# chat_docs: list of {"name": str, "text": str} — the active document pool
# chat_collection: ChromaDB collection built from chat_docs
# doc_messages: conversation history for Document Chat tab
# research_messages: conversation history for Research Agent tab
# research_results: list of {"label": str, "text": str} collected during research

if "chat_docs"          not in st.session_state:
    # Load default docs (Hassan's resume) on first run
    default_docs = load_default_docs("my_docs")
    st.session_state.chat_docs = default_docs

if "chat_collection"    not in st.session_state:
    if st.session_state.chat_docs:
        st.session_state.chat_collection = build_collection_from_docs(
            st.session_state.chat_docs, "chat_docs"
        )
    else:
        st.session_state.chat_collection = None

if "doc_messages"       not in st.session_state:
    st.session_state.doc_messages       = []

if "research_messages"  not in st.session_state:
    st.session_state.research_messages  = []

if "research_log"       not in st.session_state:
    st.session_state.research_log       = []

# research_results: accumulated text snippets from research
# each entry: {"label": str, "text": str, "added": bool}
if "research_results"   not in st.session_state:
    st.session_state.research_results   = []

# switch_to_chat: flag set when user clicks "Start Chatting" in Research tab
if "switch_to_chat"     not in st.session_state:
    st.session_state.switch_to_chat     = False

# force_doc_title: overrides title to "Document Chat" for one render cycle
# after Start Chatting is clicked, regardless of what query params say
if "force_doc_title"    not in st.session_state:
    st.session_state.force_doc_title    = False

# research_history: list of completed research runs
# each entry: {"question": str, "report": str, "steps": int, "results": list}
if "research_history"   not in st.session_state:
    st.session_state.research_history   = []


# ════════════════════════════════════════════════════════════════════════════════
# TABS  —  Document Chat is Tab 1 (default)
# Active tab is detected via query params set by a JS MutationObserver below.
# ════════════════════════════════════════════════════════════════════════════════

# Read active tab from query params (set by JS observer)
# Default to "doc" on first load
active_tab = st.query_params.get("tab", "doc")

tab1, tab2 = st.tabs(["📄 Document Chat", "🔍 Research Agent"])

# ── JS MutationObserver — watches tab clicks and syncs to query params ────────
# When the user clicks a tab, the selected tab's aria-selected attribute changes
# to "true". The observer detects this and updates the URL query param,
# which causes Streamlit to rerun and pick up the new active_tab value.
components.html(
    """
    <script>
    (function() {
        function syncTab() {
            const tabs = window.parent.document.querySelectorAll(
                'button[data-baseweb="tab"]'
            );
            if (!tabs.length) return;
            const activeIndex = Array.from(tabs).findIndex(
                t => t.getAttribute("aria-selected") === "true"
            );
            const paramValue = activeIndex === 1 ? "research" : "doc";
            const url = new URL(window.parent.location.href);
            if (url.searchParams.get("tab") !== paramValue) {
                url.searchParams.set("tab", paramValue);
                window.parent.history.replaceState({}, "", url.toString());
            }
        }
        // Observe the whole document for attribute changes on tab buttons
        const observer = new MutationObserver(syncTab);
        observer.observe(window.parent.document.body, {
            subtree: true, attributes: true, attributeFilter: ["aria-selected"]
        });
        syncTab(); // run once immediately
    })();
    </script>
    """,
    height=0
)


# ════════════════════════════════════════════════════════════════════════════════
# TAB 1 — DOCUMENT CHAT
# ════════════════════════════════════════════════════════════════════════════════

with tab1:

    # Title logic:
    # force_doc_title takes priority — set when Start Chatting is clicked.
    # It overrides the query param for one render cycle so the title is
    # correct before the MutationObserver has a chance to sync the URL.

    if st.session_state.force_doc_title or active_tab == "doc":
        st.title("📄 Document Chat")
        st.session_state.force_doc_title = False
    else:
        st.title("🧠 Learn Smart")


    # ── JS tab switch when user clicks "Start Chatting" from Research tab ──────
    # Streamlit doesn't expose a tab-switching API, so we inject a tiny JS
    # snippet that clicks the first tab button (Document Chat) in the DOM.
    if st.session_state.switch_to_chat:
        st.session_state.switch_to_chat  = False
        st.session_state.force_doc_title = True   # ensures title shows "Document Chat"
        research_doc_names = [
            d["name"].replace("research:", "")
            for d in st.session_state.chat_docs
            if d["name"].startswith("research:")
        ]
        st.success(
            f"✅ Ready to chat! {len(research_doc_names)} research document(s) loaded: "
            + ", ".join(research_doc_names)
        )
        # Click the first tab (Document Chat) via JS
        components.html(
            """
            <script>
            const tabs = window.parent.document.querySelectorAll(
                'button[data-baseweb="tab"]'
            );
            if (tabs.length > 0) { tabs[0].click(); }
            </script>
            """,
            height=0
        )

    # ── Welcome message explaining the sample doc and options ─────────────────
    # NOTE: This info box is shown when the default sample document is the only
    # document loaded. It explains what the user can do next.
    doc_names = [d["name"] for d in st.session_state.chat_docs]
    is_default_only = (
        len(st.session_state.chat_docs) > 0
        and all(
            name in ["hassan_resume.txt", "hassan_resume.md"]
            for name in doc_names
        )
    )

    if is_default_only and not st.session_state.doc_messages:
        st.info(
            "👋 **Welcome to Learn Smart!**\n\n"
            "The default document loaded is **Hassan's resume** — this is just a sample "
            "to show you how the Document Chat works.\n\n"
            "**Here's what you can do:**\n"
            "- 💬 Chat with this sample document right now\n"
            "- 📂 Upload your own document (.txt, .md, .pdf, .docx) using the sidebar\n"
            "- 🔍 Switch to the **Research Agent** tab to research any topic, then add "
            "the findings or source URLs directly to this chat"
        )
    elif not st.session_state.chat_docs:
        st.warning(
            "No documents loaded. Upload a document in the sidebar or use the "
            "Research Agent tab to add research content here."
        )

    # ── Dynamic subtitle based on active docs ─────────────────────────────────
    if doc_names:
        if is_default_only:
            st.caption("🧑‍💼 Chatting with: Hassan's resume (sample document)")
        else:
            st.caption(f"💬 Chatting with: {', '.join(doc_names)}")

    # ── Sidebar ───────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown(
        '''<div style="display:flex;align-items:center;gap:12px;padding:8px 0 14px;">
          <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 80 80" width="52" height="52" style="flex-shrink:0;">
            <rect width="80" height="80" rx="18" fill="#0f172a"/>
            <circle cx="40" cy="40" r="8" fill="#f59e0b"/>
            <circle cx="40" cy="40" r="5" fill="#fbbf24"/>
            <line x1="40" y1="32" x2="40" y2="14" stroke="#f97316" stroke-width="2" stroke-linecap="round"/>
            <line x1="40" y1="48" x2="40" y2="66" stroke="#f97316" stroke-width="2" stroke-linecap="round"/>
            <line x1="32" y1="40" x2="14" y2="40" stroke="#f97316" stroke-width="2" stroke-linecap="round"/>
            <line x1="48" y1="40" x2="66" y2="40" stroke="#f97316" stroke-width="2" stroke-linecap="round"/>
            <line x1="34" y1="34" x2="20" y2="20" stroke="#fcd34d" stroke-width="1.5" stroke-linecap="round"/>
            <line x1="46" y1="34" x2="60" y2="20" stroke="#fcd34d" stroke-width="1.5" stroke-linecap="round"/>
            <line x1="34" y1="46" x2="20" y2="60" stroke="#fcd34d" stroke-width="1.5" stroke-linecap="round"/>
            <line x1="46" y1="46" x2="60" y2="60" stroke="#fcd34d" stroke-width="1.5" stroke-linecap="round"/>
            <circle cx="40" cy="13" r="4.5" fill="#f59e0b" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="40" cy="67" r="4.5" fill="#f59e0b" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="13" cy="40" r="4.5" fill="#f59e0b" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="67" cy="40" r="4.5" fill="#f59e0b" stroke="#fbbf24" stroke-width="1.5"/>
            <circle cx="20" cy="20" r="3.5" fill="#0f172a" stroke="#fcd34d" stroke-width="1.5"/>
            <circle cx="60" cy="20" r="3.5" fill="#0f172a" stroke="#fcd34d" stroke-width="1.5"/>
            <circle cx="20" cy="60" r="3.5" fill="#0f172a" stroke="#fcd34d" stroke-width="1.5"/>
            <circle cx="60" cy="60" r="3.5" fill="#0f172a" stroke="#fcd34d" stroke-width="1.5"/>
          </svg>
          <div style="line-height:1.3;">
            <div style="font-size:1rem;font-weight:700;color:#e8f0fe;font-family:sans-serif;">Learn Smart</div>
            <div style="font-size:0.65rem;color:#fbbf24;letter-spacing:1.5px;font-family:sans-serif;">AI-POWERED</div>
          </div>
        </div>''',
        unsafe_allow_html=True
    )
        st.divider()

        # ── Files currently in chat ───────────────────────────────────────────
        st.subheader("📂 Files in Chat")

        if st.session_state.chat_docs:
            for i, doc in enumerate(st.session_state.chat_docs):
                col1, col2 = st.columns([4, 1])
                with col1:
                    # Show a small icon based on file type
                    ext  = os.path.splitext(doc["name"])[1].lower()
                    icon = {"pdf": "📕", ".docx": "📘", ".md": "📝"}.get(ext, "📄")
                    st.caption(f"{icon} {doc['name']}")
                with col2:
                    if st.button("✕", key=f"remove_{i}", help=f"Remove {doc['name']}"):
                        st.session_state.chat_docs.pop(i)
                        st.session_state.doc_messages = []  # reset chat
                        rebuild_chat_collection()
                        st.rerun()
        else:
            st.caption("No documents loaded.")

        st.divider()

        # ── Upload your own document ──────────────────────────────────────────
        st.subheader("Upload a Document")
        st.info(
            "**Accepted formats:**\n"
            "- `.txt` — plain text\n"
            "- `.md` — markdown\n"
            "- `.pdf` — PDF (must have selectable text, not scanned)\n"
            "- `.docx` — Word document\n\n"
            "Uploading a file **replaces** the default resume in chat."
        )

        uploaded_file = st.file_uploader(
            "Choose a file",
            type=["txt", "md", "pdf", "docx"],
            help="Supported: .txt  .md  .pdf  .docx"
        )

        if uploaded_file:
            already_loaded = any(
                d["name"] == uploaded_file.name
                and not d["name"].startswith("research:")
                for d in st.session_state.chat_docs
            )
            if already_loaded:
                st.warning(f"\'{uploaded_file.name}\' is already in the chat.")
            else:
                if st.button("Add to chat", type="primary", use_container_width=True):
                    with st.spinner("Extracting and indexing..."):
                        text = extract_text_from_file(uploaded_file)
                    if text.strip():
                        # Remove default (resume) docs — keep only research docs
                        # already added, then append the new upload
                        st.session_state.chat_docs = [
                            d for d in st.session_state.chat_docs
                            if d["name"].startswith("research:")
                        ]
                        st.session_state.chat_docs.append({
                            "name": uploaded_file.name,
                            "text": text
                        })
                        st.session_state.doc_messages = []
                        rebuild_chat_collection()
                        st.success(f"Loaded: {uploaded_file.name}")
                        st.rerun()
                    else:
                        st.error(
                            "Could not extract text. Check that your file:\n"
                            "- Is not password protected\n"
                            "- Contains selectable text (not a scanned image)\n"
                            "- Is not corrupted"
                        )

        st.divider()

        # ── Controls ──────────────────────────────────────────────────────────
        if st.button("🗑️ Clear conversation", use_container_width=True, key="clear_doc"):
            st.session_state.doc_messages = []
            st.rerun()

        if st.button("🔄 Reset to default document", use_container_width=True):
            default_docs = load_default_docs("my_docs")
            st.session_state.chat_docs    = default_docs
            st.session_state.doc_messages = []
            rebuild_chat_collection()
            st.rerun()

        st.divider()
        show_context = st.toggle("Show retrieved context", value=False)
        st.caption("See which document sections were used for each answer.")

    # ── Render messages ───────────────────────────────────────────────────────
    for message in st.session_state.doc_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # ── Chat input ────────────────────────────────────────────────────────────
    chat_disabled    = st.session_state.chat_collection is None
    chat_placeholder = (
        "Ask a question about your documents..."
        if not chat_disabled
        else "Add a document above to start chatting..."
    )

    if prompt := st.chat_input(chat_placeholder, key="doc_input", disabled=chat_disabled):
        with st.chat_message("user"):
            st.markdown(prompt)
        st.session_state.doc_messages.append({"role": "user", "content": prompt})

        context = retrieve_context(st.session_state.chat_collection, prompt)

        if show_context:
            with st.expander("📄 Retrieved document sections", expanded=False):
                st.text(context)

        doc_label = ", ".join(doc_names) if doc_names else "the documents"
        augmented = f"""Here is relevant information from {doc_label}:

{context}

---
Question: {prompt}"""

        history = []
        for msg in st.session_state.doc_messages[:-1]:
            history.append({"role": msg["role"], "content": msg["content"]})
        history.append({"role": "user", "content": augmented})

        active_system_prompt = build_chat_system_prompt(doc_names)

        with st.chat_message("assistant"):
            placeholder   = st.empty()
            full_response = ""
            with client.messages.stream(
                model="claude-sonnet-4-6",
                max_tokens=1024,
                system=active_system_prompt,
                messages=history
            ) as stream:
                for chunk in stream.text_stream:
                    full_response += chunk
                    placeholder.markdown(full_response + "▌")
            placeholder.markdown(full_response)

        st.session_state.doc_messages.append({
            "role": "assistant", "content": full_response
        })
        chunk_count = len(context.split("[Document")) - 1
        st.caption(f"Retrieved {chunk_count} document sections")


# ════════════════════════════════════════════════════════════════════════════════
# TAB 2 — RESEARCH AGENT
# ════════════════════════════════════════════════════════════════════════════════

with tab2:
    st.title("🔍 Research Agent")
    st.caption(
        "Research any topic. The agent searches the web and compiles a report. "
        "You can then **add the report or individual sources to Document Chat** for deeper exploration."
    )

    # ── How it works ──────────────────────────────────────────────────────────
    with st.expander("⚙️ How the research agent works", expanded=False):
        st.markdown("""
The research agent uses **agentic looping** — Claude actively searches the web and
reads pages before responding, rather than answering from memory alone.

**Step by step:**
1. You submit a research question
2. Claude decides the best search queries to answer it
3. It calls `web_search` and reads the results
4. If a result looks useful, it calls `fetch_page` to read that page in full
5. It repeats steps 3–4 as needed, then compiles a structured report
6. The loop stops when Claude has enough information

**Adding research to Document Chat:**
Each research result and source URL appears below with an **Add to Chat** button.
Click it to send that content to the Document Chat tab so you can ask follow-up
questions grounded in the exact sources the agent found.

**What you'll see in the activity log:**
Each tool call is shown in real time — watch Claude search and reason step by step.
        """)

    with st.expander("💡 Tips for getting the best research results", expanded=False):
        st.markdown("""
**Be specific, not broad**
- ❌ `"Tell me about AI"` — too vague, results will be shallow
- ✅ `"What are the most practical AI agent use cases for small businesses in 2025?"` — focused

**Include context**
- ❌ `"What is RAG?"`
- ✅ `"What is RAG (retrieval-augmented generation) and how is it used in enterprise chatbots?"`

**Ask for comparisons**
- `"Compare LangChain vs LlamaIndex for RAG pipelines — pros, cons, when to use each"`

**Ask for actionable outputs**
- `"What are the top 5 steps a freelance developer should take to land their first AI agent project?"`

**Specify your audience**
- `"Explain prompt engineering to a non-technical business owner"`

**What the agent is NOT good at:**
- Very recent breaking news (search index may not have it yet)
- Pages behind a login or paywall
- Exact statistics — always verify numbers from the original source
        """)

    st.divider()

    # ── Documents added to chat from research ─────────────────────────────────
    research_added = [
        d["name"] for d in st.session_state.chat_docs
        if d["name"].startswith("research:")
    ]

    if research_added:
        st.success(
            f"✅ {len(research_added)} research item(s) added to Document Chat: "
            + ", ".join(r.replace("research:", "") for r in research_added)
        )
        # ── Start Chatting button — shown once at least one item is added ─────
        if st.button(
            "💬 Start Chatting →",
            type="primary",
            use_container_width=True,
            key="start_chatting"
        ):
            # Switch to Tab 1 by setting a flag and rerunning.
            # Streamlit doesn't support programmatic tab switching directly,
            # so we use a query param to signal the user to click Tab 1.
            st.session_state.switch_to_chat = True
            st.rerun()

    # ── Selectable research results ───────────────────────────────────────────
    if st.session_state.research_results:
        st.subheader("📋 Research Results — Add to Chat")
        st.caption(
            "Click **Add to Chat** on any item below to include it in the "
            "Document Chat tab. The default sample resume will be removed automatically."
        )

        for i, result in enumerate(st.session_state.research_results):
            already = any(
                d["name"] == f"research:{result['label']}"
                for d in st.session_state.chat_docs
            )
            with st.container():
                col1, col2 = st.columns([5, 1])
                with col1:
                    with st.expander(
                        f"{'✅' if already else '📄'} {result['label']}",
                        expanded=False
                    ):
                        st.text(result["text"][:800] + ("..." if len(result["text"]) > 800 else ""))
                with col2:
                    if already:
                        st.caption("Added ✓")
                    else:
                        if st.button(
                            "Add to Chat",
                            key=f"add_result_{i}",
                            use_container_width=True
                        ):
                            # Remove default (resume) docs before adding research content
                            # so the chat is focused only on the selected research material
                            st.session_state.chat_docs = [
                                d for d in st.session_state.chat_docs
                                if d["name"].startswith("research:")
                            ]
                            st.session_state.chat_docs.append({
                                "name": f"research:{result['label']}",
                                "text": result["text"]
                            })
                            st.session_state.doc_messages = []
                            rebuild_chat_collection()
                            st.rerun()

        st.divider()

    # ── Current session activity log ─────────────────────────────────────────
    if st.session_state.research_log:
        with st.expander(
            f"📋 Activity log ({len(st.session_state.research_log)} steps)",
            expanded=False
        ):
            for entry in st.session_state.research_log:
                st.caption(entry)

    # ── Render current research messages ──────────────────────────────────────
    for message in st.session_state.research_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # ── Clear current research button ─────────────────────────────────────────
    if st.session_state.research_messages:
        if st.button("🗑️ Clear current research", key="clear_research"):
            st.session_state.research_messages = []
            st.session_state.research_log      = []
            st.session_state.research_results  = []
            st.rerun()

    st.divider()

    # ── Research history ──────────────────────────────────────────────────────
    if st.session_state.research_history:
        st.subheader(f"🕘 Research History ({len(st.session_state.research_history)} sessions)")
        st.caption("Past research sessions — expand any to read the report or add sources to Document Chat.")

        # Show newest first
        for h_idx, entry in enumerate(reversed(st.session_state.research_history)):
            real_idx = len(st.session_state.research_history) - 1 - h_idx
            with st.expander(
                f"{'🔍'} {entry['question'][:70]}{'...' if len(entry['question']) > 70 else ''}  "
                f"·  {entry['steps']} steps",
                expanded=False
            ):
                st.markdown(entry["report"])
                st.divider()
                st.caption("**Add sources from this session to Document Chat:**")
                for r_idx, result in enumerate(entry["results"]):
                    already = any(
                        d["name"] == f"research:{result['label']}"
                        for d in st.session_state.chat_docs
                    )
                    col1, col2 = st.columns([5, 1])
                    with col1:
                        st.caption(f"{'✅' if already else '📄'} {result['label'][:70]}")
                    with col2:
                        if already:
                            st.caption("Added ✓")
                        else:
                            if st.button(
                                "Add",
                                key=f"hist_{real_idx}_{r_idx}",
                                use_container_width=True
                            ):
                                st.session_state.chat_docs = [
                                    d for d in st.session_state.chat_docs
                                    if d["name"].startswith("research:")
                                ]
                                st.session_state.chat_docs.append({
                                    "name": f"research:{result['label']}",
                                    "text": result["text"]
                                })
                                st.session_state.doc_messages = []
                                rebuild_chat_collection()
                                st.rerun()

        if st.button("🗑️ Clear all history", key="clear_history"):
            st.session_state.research_history = []
            st.rerun()

    # ── Research input ────────────────────────────────────────────────────────
    if research_prompt := st.chat_input("Enter a research question...", key="research_input"):

        with st.chat_message("user"):
            st.markdown(research_prompt)
        st.session_state.research_messages.append({"role": "user", "content": research_prompt})
        st.session_state.research_log     = []
        st.session_state.research_results = []

        messages         = [{"role": "user", "content": research_prompt}]
        step             = 1
        status_container = st.empty()

        # Track URLs and text fetched during this research run
        fetched_pages  = {}  # url → text
        search_results = {}  # query → text

        with st.chat_message("assistant"):
            response_placeholder = st.empty()

            while True:
                response = client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=4096,
                    system=RESEARCH_SYSTEM_PROMPT,
                    tools=research_tools,
                    messages=messages
                )

                if response.stop_reason == "end_turn":
                    final_text = ""
                    for block in response.content:
                        if hasattr(block, "text"):
                            final_text = block.text
                            break

                    status_container.empty()
                    response_placeholder.markdown(final_text)
                    st.session_state.research_messages.append({
                        "role": "assistant", "content": final_text
                    })

                    # ── Save selectable results ────────────────────────────────
                    # Add the full report as a selectable item
                    run_results = [{
                        "label": f"Full report — {research_prompt[:50]}",
                        "text":  final_text
                    }]
                    st.session_state.research_results.append(run_results[0])
                    # Add each fetched page as a selectable item
                    for url, text in fetched_pages.items():
                        entry = {"label": url[:60], "text": text}
                        st.session_state.research_results.append(entry)
                        run_results.append(entry)

                    # ── Save to research history ───────────────────────────────
                    st.session_state.research_history.append({
                        "question": research_prompt,
                        "report":   final_text,
                        "steps":    step - 1,
                        "results":  run_results
                    })
                    break

                if response.stop_reason == "tool_use":
                    messages.append({"role": "assistant", "content": response.content})
                    tool_results = []

                    for block in response.content:
                        if block.type == "tool_use":
                            label = (
                                f"🔎 Searching: *{block.input.get('query', '')}*"
                                if block.name == "web_search"
                                else f"📖 Reading: *{block.input.get('url', '')[:60]}...*"
                            )
                            status_container.info(f"Step {step}: {label}")
                            log_entry = f"Step {step} — {block.name}: {list(block.input.values())[0][:80]}"
                            st.session_state.research_log.append(log_entry)

                            result = run_research_tool(block.name, block.input)

                            # Store fetched page text for "Add to Chat"
                            if block.name == "fetch_page":
                                fetched_pages[block.input.get("url", f"page_{step}")] = result

                            tool_results.append({
                                "type": "tool_result",
                                "tool_use_id": block.id,
                                "content": result
                            })
                            step += 1

                    messages.append({"role": "user", "content": tool_results})

                else:
                    status_container.empty()
                    break

        st.caption(f"Research completed in {step - 1} steps")
        st.rerun()  # refresh to show the Add to Chat buttons immediately
