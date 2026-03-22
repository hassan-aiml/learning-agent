import os
import io
import requests
import chromadb
import streamlit as st
from datetime import datetime
from dotenv import load_dotenv
from anthropic import Anthropic
from bs4 import BeautifulSoup
from sentence_transformers import SentenceTransformer

load_dotenv()

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AI Assistant",
    page_icon="🤖",
    layout="centered"
)


# ── Cached resources ──────────────────────────────────────────────────────────
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

@st.cache_resource
def load_claude_client():
    return Anthropic()

model  = load_embedding_model()
client = load_claude_client()


# ════════════════════════════════════════════════════════════════════════════════
# SHARED UTILITIES
# ════════════════════════════════════════════════════════════════════════════════

def chunk_text(text: str, chunk_size: int = 300, overlap: int = 50) -> list:
    """Split text into overlapping chunks for RAG retrieval."""
    words  = text.split()
    chunks = []
    start  = 0
    while start < len(words):
        chunk = " ".join(words[start:start + chunk_size])
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def extract_text(uploaded_file) -> str:
    """
    Extract plain text from uploaded files.
    Supports: .txt, .md, .pdf, .docx
    Returns extracted text as a string.
    """
    filename = uploaded_file.name.lower()

    # Plain text and markdown
    if filename.endswith((".txt", ".md")):
        return uploaded_file.read().decode("utf-8")

    # PDF — uses PyMuPDF (fitz)
    elif filename.endswith(".pdf"):
        try:
            import fitz  # PyMuPDF
            pdf_bytes = uploaded_file.read()
            doc       = fitz.open(stream=pdf_bytes, filetype="pdf")
            pages     = [page.get_text() for page in doc]
            return "\n\n".join(pages)
        except ImportError:
            st.error("PyMuPDF not installed. Run: pip install pymupdf")
            return ""
        except Exception as e:
            st.error(f"Could not read PDF: {e}")
            return ""

    # Word document — uses python-docx
    elif filename.endswith(".docx"):
        try:
            from docx import Document
            doc_bytes = io.BytesIO(uploaded_file.read())
            doc       = Document(doc_bytes)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            st.error("python-docx not installed. Run: pip install python-docx")
            return ""
        except Exception as e:
            st.error(f"Could not read Word document: {e}")
            return ""

    else:
        st.error(f"Unsupported file type: {filename}")
        return ""


def title_from_filename(filename: str) -> str:
    """'company_faq.pdf' → '💬 Chat with: company faq'"""
    name = os.path.splitext(filename)[0]
    name = name.replace("_", " ").replace("-", " ")
    return f"💬 Chat with: {name}"


def retrieve_context(collection, query: str, top_k: int = 4) -> str:
    """Semantic search over the vector DB, returns formatted chunks."""
    query_embedding = model.encode([query]).tolist()
    results = collection.query(query_embeddings=query_embedding, n_results=top_k)
    chunks  = results["documents"][0]
    return "\n\n".join([f"[Document {i+1}]: {chunk}" for i, chunk in enumerate(chunks)])


# ════════════════════════════════════════════════════════════════════════════════
# TAB 1 — DOCUMENT CHAT
# ════════════════════════════════════════════════════════════════════════════════

# ── Build KB from my_docs/ folder (Hassan default) ───────────────────────────
@st.cache_resource
def build_default_kb(folder: str = "my_docs") -> tuple:
    """Loads Hassan's resume from my_docs/ at startup. Cached for the session."""
    if not os.path.exists(folder):
        return None, 0, []

    files = [f for f in os.listdir(folder) if f.endswith((".txt", ".md", ".pdf", ".docx"))]
    if not files:
        return None, 0, []

    all_ids   = []
    all_texts = []
    filenames = []

    for filename in files:
        filepath = os.path.join(folder, filename)
        ext      = filename.lower()

        if ext.endswith((".txt", ".md")):
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read().strip()
        elif ext.endswith(".pdf"):
            try:
                import fitz
                doc  = fitz.open(filepath)
                text = "\n\n".join([page.get_text() for page in doc])
            except Exception:
                continue
        elif ext.endswith(".docx"):
            try:
                from docx import Document
                doc  = Document(filepath)
                text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            except Exception:
                continue
        else:
            continue

        if not text.strip():
            continue

        chunks = chunk_text(text)
        for i, chunk in enumerate(chunks):
            all_ids.append(f"{filename}_chunk_{i}")
            all_texts.append(chunk)
        filenames.append(filename)

    if not all_texts:
        return None, 0, []

    chroma_client = chromadb.PersistentClient(path="./knowledge_base")
    try:
        chroma_client.delete_collection("default_docs")
    except:
        pass
    collection     = chroma_client.create_collection("default_docs")
    all_embeddings = model.encode(all_texts, show_progress_bar=False).tolist()
    collection.add(ids=all_ids, embeddings=all_embeddings, documents=all_texts)

    return collection, len(all_texts), filenames


def build_uploaded_kb(uploaded_file) -> tuple:
    """Builds a fresh KB from a user-uploaded file."""
    text = extract_text(uploaded_file)
    if not text.strip():
        return None, 0

    chunks         = chunk_text(text)
    all_ids        = [f"upload_chunk_{i}" for i in range(len(chunks))]
    all_embeddings = model.encode(chunks, show_progress_bar=False).tolist()

    chroma_client = chromadb.PersistentClient(path="./knowledge_base")
    try:
        chroma_client.delete_collection("uploaded_docs")
    except:
        pass
    collection = chroma_client.create_collection("uploaded_docs")
    collection.add(ids=all_ids, embeddings=all_embeddings, documents=chunks)

    return collection, len(chunks)


# ── System prompts ────────────────────────────────────────────────────────────
HASSAN_SYSTEM_PROMPT = """You are a professional HR assistant helping recruiters and
hiring managers learn about Hassan M. Hai from his resume.

IMPORTANT RULES:
- Only answer using the resume context provided — never invent experience or skills
- If the resume doesn't contain enough information to answer, say so clearly
- Refer to the candidate by first name: Hassan
- Be objective, professional, and concise
- When quoting specific facts (years, companies, technologies), be precise
- Do not speculate about things not mentioned in the resume"""

def uploaded_system_prompt(filename: str) -> str:
    name = os.path.splitext(filename)[0].replace("_", " ").replace("-", " ")
    return f"""You are a helpful assistant answering questions about the document: '{name}'.

IMPORTANT RULES:
- Only answer using the document context provided — never invent information
- If the document doesn't contain enough information to answer, say so clearly
- Be concise, accurate, and helpful
- When quoting specific facts, be precise
- Do not speculate about things not mentioned in the document"""


# ════════════════════════════════════════════════════════════════════════════════
# TAB 2 — RESEARCH AGENT
# ════════════════════════════════════════════════════════════════════════════════

RESEARCH_SYSTEM_PROMPT = """You are a professional research assistant. Research topics 
thoroughly and produce clear, structured reports.

Your process for every research request:
1. Search for the topic to get an overview
2. If you find useful URLs, fetch 1-2 of them for more detail
3. Synthesise everything into a structured report

Your report format:
# [Topic Title]
## Summary
(2-3 sentence overview)
## Key Findings
(bullet points of the most important facts)
## Details
(more depth on the most interesting points)
## Sources
(list any URLs you found)

Be factual and concise. Only report what you found — never make up information."""

research_tools = [
    {
        "name": "web_search",
        "description": "Search the web for information on a topic. Use this first.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Specific search query"}
            },
            "required": ["query"]
        }
    },
    {
        "name": "fetch_page",
        "description": "Read the full content of a webpage URL for more detail.",
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "Full URL including https://"}
            },
            "required": ["url"]
        }
    }
]

def web_search(query: str) -> str:
    try:
        url    = "https://api.duckduckgo.com/"
        params = {"q": query, "format": "json", "no_html": 1, "skip_disambig": 1}
        r      = requests.get(url, params=params, timeout=10)
        data   = r.json()
        results = []
        if data.get("AbstractText"):
            results.append(f"Summary: {data['AbstractText']}")
            results.append(f"Source: {data.get('AbstractURL', '')}")
        for topic in data.get("RelatedTopics", [])[:5]:
            if isinstance(topic, dict) and topic.get("Text"):
                results.append(f"- {topic['Text']}")
        return "\n".join(results) if results else f"No direct results for '{query}'."
    except Exception as e:
        return f"Search error: {str(e)}"

def fetch_page(url: str) -> str:
    try:
        headers  = {"User-Agent": "Mozilla/5.0 (research bot)"}
        r        = requests.get(url, headers=headers, timeout=10)
        r.raise_for_status()
        soup     = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text  = soup.get_text(separator="\n", strip=True)
        lines = [l for l in text.splitlines() if len(l.strip()) > 40]
        return "\n".join(lines[:80]) or "No readable content found."
    except requests.exceptions.Timeout:
        return "Error: Page timed out."
    except requests.exceptions.HTTPError as e:
        return f"Error: HTTP {e.response.status_code}"
    except Exception as e:
        return f"Fetch error: {str(e)}"

def run_research_tool(name: str, inputs: dict) -> str:
    try:
        if name == "web_search": return web_search(inputs["query"])
        if name == "fetch_page": return fetch_page(inputs["url"])
        return f"Unknown tool: {name}"
    except Exception as e:
        return f"Tool '{name}' failed: {str(e)}"


# ════════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ════════════════════════════════════════════════════════════════════════════════

# Document chat state
if "doc_messages"           not in st.session_state:
    st.session_state.doc_messages           = []
if "uploaded_collection"    not in st.session_state:
    st.session_state.uploaded_collection    = None
if "uploaded_filename"      not in st.session_state:
    st.session_state.uploaded_filename      = None
if "uploaded_chunks"        not in st.session_state:
    st.session_state.uploaded_chunks        = 0

# Research agent state
if "research_messages"      not in st.session_state:
    st.session_state.research_messages      = []
if "research_log"           not in st.session_state:
    st.session_state.research_log           = []


# ════════════════════════════════════════════════════════════════════════════════
# LOAD DEFAULT KB
# ════════════════════════════════════════════════════════════════════════════════

default_collection, default_chunks, default_files = build_default_kb("my_docs")


# ════════════════════════════════════════════════════════════════════════════════
# TABS
# ════════════════════════════════════════════════════════════════════════════════

tab1, tab2 = st.tabs(["📄 Document Chat", "🔍 Research Agent"])


# ════════════════════════════════════════════════════════════════════════════════
# TAB 1 — DOCUMENT CHAT
# ════════════════════════════════════════════════════════════════════════════════

with tab1:

    # ── Active mode ───────────────────────────────────────────────────────────
    using_upload      = st.session_state.uploaded_collection is not None
    active_collection = st.session_state.uploaded_collection if using_upload else default_collection
    active_prompt     = (
        uploaded_system_prompt(st.session_state.uploaded_filename)
        if using_upload else HASSAN_SYSTEM_PROMPT
    )

    # ── Dynamic header ────────────────────────────────────────────────────────
    if using_upload:
        st.title(title_from_filename(st.session_state.uploaded_filename))
        st.caption(f"Answering questions from: **{st.session_state.uploaded_filename}**")
    else:
        st.title("🧑‍💼 Learn about Hassan")
        st.caption("Ask anything about Hassan's experience, skills, or background.")

    # ── Sidebar controls for Tab 1 ────────────────────────────────────────────
    with st.sidebar:
        st.header("📄 Document Chat")

        # Default doc status
        if default_collection:
            st.success("Hassan's resume loaded")
            st.metric("Default chunks", default_chunks)
        else:
            st.warning("No default document in `my_docs/`")

        st.divider()
        st.subheader("Upload Your Own Document")
        st.info(
            "**Accepted formats:**\n"
            "- `.txt` — plain text\n"
            "- `.md` — markdown\n"
            "- `.pdf` — PDF document\n"
            "- `.docx` — Word document\n\n"
            "**Note:** Make sure your PDF or Word file contains selectable text "
            "(not a scanned image). Any document type works: resumes, manuals, "
            "FAQs, reports, or notes."
        )

        uploaded_file = st.file_uploader(
            "Choose a file",
            type=["txt", "md", "pdf", "docx"],
            help="Supported: .txt  .md  .pdf  .docx"
        )

        if uploaded_file:
            if st.button("Load document", type="primary", use_container_width=True):
                with st.spinner("Extracting and indexing your document..."):
                    collection, total = build_uploaded_kb(uploaded_file)
                if collection:
                    st.session_state.uploaded_collection = collection
                    st.session_state.uploaded_filename   = uploaded_file.name
                    st.session_state.uploaded_chunks     = total
                    st.session_state.doc_messages        = []
                    st.success(f"Loaded! {total} chunks indexed.")
                    st.rerun()
                else:
                    st.error(
                        "Could not extract text. Check that your file:\n"
                        "- Is not password protected\n"
                        "- Contains selectable text (not a scanned image)\n"
                        "- Is not corrupted"
                    )

        if using_upload:
            st.divider()
            st.caption(f"Active: **{st.session_state.uploaded_filename}**")
            st.metric("Upload chunks", st.session_state.uploaded_chunks)
            if st.button("↩️ Back to Hassan's resume", use_container_width=True):
                st.session_state.uploaded_collection = None
                st.session_state.uploaded_filename   = None
                st.session_state.uploaded_chunks     = 0
                st.session_state.doc_messages        = []
                st.rerun()

        st.divider()
        if st.button("🗑️ Clear chat", use_container_width=True, key="clear_doc"):
            st.session_state.doc_messages = []
            st.rerun()

        show_context = st.toggle("Show retrieved context", value=False)
        st.caption("See which document sections were used for each answer.")

    # ── Guard ─────────────────────────────────────────────────────────────────
    if not active_collection:
        st.error("No document loaded.")
        st.info("Add a file to `my_docs/` or upload one in the sidebar.")
        st.stop()

    # ── Render messages ───────────────────────────────────────────────────────
    for message in st.session_state.doc_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # ── Chat input ────────────────────────────────────────────────────────────
    chat_placeholder = (
        f"Ask about {os.path.splitext(st.session_state.uploaded_filename)[0]}..."
        if using_upload else "Ask about Hassan..."
    )

    if prompt := st.chat_input(chat_placeholder, key="doc_input"):
        with st.chat_message("user"):
            st.markdown(prompt)
        st.session_state.doc_messages.append({"role": "user", "content": prompt})

        context = retrieve_context(active_collection, prompt)

        if show_context:
            with st.expander("📄 Retrieved document sections", expanded=False):
                st.text(context)

        doc_label = st.session_state.uploaded_filename if using_upload else "Hassan's resume"
        augmented = f"""Here is relevant information from {doc_label}:

{context}

---
Question: {prompt}"""

        history = []
        for msg in st.session_state.doc_messages[:-1]:
            history.append({"role": msg["role"], "content": msg["content"]})
        history.append({"role": "user", "content": augmented})

        with st.chat_message("assistant"):
            placeholder   = st.empty()
            full_response = ""
            with client.messages.stream(
                model="claude-sonnet-4-6",
                max_tokens=1024,
                system=active_prompt,
                messages=history
            ) as stream:
                for chunk in stream.text_stream:
                    full_response += chunk
                    placeholder.markdown(full_response + "▌")
            placeholder.markdown(full_response)

        st.session_state.doc_messages.append({"role": "assistant", "content": full_response})
        chunk_count = len(context.split("[Document")) - 1
        st.caption(f"Retrieved {chunk_count} document sections")


# ════════════════════════════════════════════════════════════════════════════════
# TAB 2 — RESEARCH AGENT
# ════════════════════════════════════════════════════════════════════════════════

with tab2:
    st.title("🔍 Research Agent")
    st.caption("Ask any research question and the agent will search the web and compile a structured report.")

    # ── How it works ──────────────────────────────────────────────────────────
    with st.expander("⚙️ How the research agent works", expanded=False):
        st.markdown("""
The research agent uses a technique called **agentic looping** — Claude doesn't just answer 
from memory, it actively searches the web and reads pages before responding.

**Step by step:**
1. You submit a research question
2. Claude decides which search queries will best answer it
3. It calls the `web_search` tool and reads the results
4. If a result looks useful, it calls `fetch_page` to read that page in full
5. It repeats steps 3–4 as needed, then compiles everything into a structured report
6. The loop only stops when Claude decides it has enough information

**What you'll see in the activity log:**
Each tool call is shown in real time — you can watch Claude search, fetch, and reason step by step.
        """)

    with st.expander("💡 Tips for getting the best research results", expanded=False):
        st.markdown("""
**Be specific, not broad**
- ❌ `"Tell me about AI"` — too vague, results will be shallow
- ✅ `"What are the most effective use cases of AI agents for small businesses in 2025?"` — focused, actionable

**Include context**
- ❌ `"What is RAG?"`
- ✅ `"What is RAG (retrieval-augmented generation) and how is it used in enterprise chatbots?"`

**Ask for comparisons**
- `"Compare LangChain vs LlamaIndex for building RAG pipelines — pros, cons, and when to use each"`

**Ask for actionable outputs**
- `"What are the top 5 steps a freelance developer should take to land their first AI agent project?"`

**Specify your audience**
- `"Explain prompt engineering to a non-technical business owner — what they need to know and why it matters"`

**What the agent is NOT good at:**
- Very recent breaking news (search index may not have it yet)
- Pages behind a login or paywall (it cannot access these)
- Exact numbers or statistics — always verify these from the original source
        """)

    st.divider()

    # ── Activity log ──────────────────────────────────────────────────────────
    if st.session_state.research_log:
        with st.expander(f"📋 Activity log ({len(st.session_state.research_log)} steps)", expanded=False):
            for entry in st.session_state.research_log:
                st.caption(entry)

    # ── Render research messages ───────────────────────────────────────────────
    for message in st.session_state.research_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # ── Clear button ──────────────────────────────────────────────────────────
    if st.session_state.research_messages:
        if st.button("🗑️ Clear research", use_container_width=False, key="clear_research"):
            st.session_state.research_messages = []
            st.session_state.research_log      = []
            st.rerun()

    # ── Research input ────────────────────────────────────────────────────────
    if research_prompt := st.chat_input(
        "Enter a research question...", key="research_input"
    ):
        with st.chat_message("user"):
            st.markdown(research_prompt)
        st.session_state.research_messages.append({"role": "user", "content": research_prompt})
        st.session_state.research_log = []

        messages = [{"role": "user", "content": research_prompt}]
        step     = 1

        # ── Status container shows live progress ──────────────────────────────
        status_container = st.empty()

        with st.chat_message("assistant"):
            response_placeholder = st.empty()

            # ── Agentic loop ──────────────────────────────────────────────────
            while True:
                response = client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=4096,
                    system=RESEARCH_SYSTEM_PROMPT,
                    tools=research_tools,
                    messages=messages
                )

                if response.stop_reason == "end_turn":
                    # Extract final text response
                    final_text = ""
                    for block in response.content:
                        if hasattr(block, "text"):
                            final_text = block.text
                            break
                    status_container.empty()
                    response_placeholder.markdown(final_text)
                    st.session_state.research_messages.append({
                        "role": "assistant", "content": final_text
                    })
                    break

                if response.stop_reason == "tool_use":
                    messages.append({"role": "assistant", "content": response.content})
                    tool_results = []

                    for block in response.content:
                        if block.type == "tool_use":
                            # Show live status
                            label = (
                                f"🔎 Searching: *{block.input.get('query', '')}*"
                                if block.name == "web_search"
                                else f"📖 Reading: *{block.input.get('url', '')[:60]}...*"
                            )
                            status_container.info(f"Step {step}: {label}")

                            # Log it
                            log_entry = f"Step {step} — {block.name}: {list(block.input.values())[0][:80]}"
                            st.session_state.research_log.append(log_entry)

                            # Run the tool
                            result = run_research_tool(block.name, block.input)
                            tool_results.append({
                                "type": "tool_result",
                                "tool_use_id": block.id,
                                "content": result
                            })
                            step += 1

                    messages.append({"role": "user", "content": tool_results})

                else:
                    status_container.empty()
                    break

        st.caption(f"Research completed in {step - 1} steps")
